"""
Export Tiptap JSON to Word (.docx) format.

Walks the Tiptap JSON document tree and maps nodes to python-docx elements.
Supports format presets for legal documents (court briefs, cover letters, declarations).
"""

import io
import re
from html import escape

from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


FORMAT_PRESETS = {
    "court_brief": {
        "font_name": "Times New Roman",
        "font_size": 12,
        "line_spacing": WD_LINE_SPACING.DOUBLE,
        "margin_inches": 1.0,
        "paragraph_alignment": WD_ALIGN_PARAGRAPH.LEFT,
    },
    "cover_letter": {
        "font_name": "Times New Roman",
        "font_size": 12,
        "line_spacing": WD_LINE_SPACING.SINGLE,
        "margin_inches": 1.0,
        "paragraph_alignment": WD_ALIGN_PARAGRAPH.LEFT,
    },
    "declaration": {
        "font_name": "Times New Roman",
        "font_size": 12,
        "line_spacing": WD_LINE_SPACING.DOUBLE,
        "margin_inches": 1.0,
        "paragraph_alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "numbered_paragraphs": True,
    },
}

DATE_PATTERN = re.compile(r"^\s*(\[[Dd]ate\]|[A-Z][a-z]+ \d{1,2}, \d{4}|\d{1,2}/\d{1,2}/\d{2,4})")


def tiptap_to_docx(tiptap_json, title="Document", export_format="court_brief", document_metadata=None):
    """Convert Tiptap JSON content to a .docx file buffer."""
    return _tiptap_to_docx_generic(
        tiptap_json,
        title=title,
        export_format=export_format,
        document_metadata=document_metadata,
    )


def tiptap_to_docx_with_template(
    tiptap_json,
    title="Document",
    export_format="court_brief",
    template_path="",
    document_metadata=None,
):
    return _tiptap_to_docx_generic(
        tiptap_json,
        title=title,
        export_format=export_format,
        template_path=template_path,
        document_metadata=document_metadata,
    )


def tiptap_to_docx_with_style_anchor(
    tiptap_json,
    title="Document",
    export_format="court_brief",
    style_anchor=None,
    document_metadata=None,
):
    if export_format == "cover_letter" and style_anchor:
        return _tiptap_to_cover_letter_docx(
            tiptap_json=tiptap_json,
            title=title,
            style_anchor=style_anchor,
        )
    return _tiptap_to_docx_generic(
        tiptap_json,
        title=title,
        export_format=export_format,
        document_metadata=document_metadata,
    )


def _tiptap_to_docx_generic(
    tiptap_json,
    title="Document",
    export_format="court_brief",
    template_path="",
    document_metadata=None,
):
    """Convert Tiptap JSON content to a .docx file buffer."""
    doc = DocxDocument(template_path) if template_path else DocxDocument()
    preset = FORMAT_PRESETS.get(export_format, FORMAT_PRESETS["court_brief"])
    preserve_template_styles = bool(template_path)

    if template_path:
        _clear_document_body(doc)
    else:
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(preset["margin_inches"])
            section.bottom_margin = Inches(preset["margin_inches"])
            section.left_margin = Inches(preset["margin_inches"])
            section.right_margin = Inches(preset["margin_inches"])
        _apply_document_page_setup(doc, document_metadata or {})

        # Configure default style
        style = doc.styles["Normal"]
        font = style.font
        font.name = preset["font_name"]
        font.size = Pt(preset["font_size"])
        pf = style.paragraph_format
        pf.line_spacing_rule = preset["line_spacing"]
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)

        # Configure heading styles
        for level in range(1, 4):
            style_name = f"Heading {level}"
            if style_name in doc.styles:
                hs = doc.styles[style_name]
                hs.font.name = preset["font_name"]
                hs.font.bold = True
                if level == 1:
                    hs.font.size = Pt(14)
                elif level == 2:
                    hs.font.size = Pt(13)
                else:
                    hs.font.size = Pt(12)
                hs.paragraph_format.space_before = Pt(12)
                hs.paragraph_format.space_after = Pt(6)

    content = tiptap_json.get("content", []) if isinstance(tiptap_json, dict) else []
    para_counter = [0]  # mutable counter for numbered paragraphs
    footnotes = []

    for node in content:
        _process_node(
            doc,
            node,
            preset,
            para_counter,
            footnotes,
            preserve_template_styles=preserve_template_styles,
        )

    if footnotes:
        doc.add_page_break()
        heading = doc.add_heading(level=2)
        heading.add_run("Footnotes")
        for footnote in footnotes:
            p = doc.add_paragraph()
            marker = p.add_run(f"[{footnote['number']}] ")
            marker.font.superscript = True
            p.add_run(footnote["text"])

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _tiptap_to_cover_letter_docx(*, tiptap_json, title, style_anchor):
    template_doc = DocxDocument(style_anchor.path)
    sample_doc = DocxDocument(style_anchor.path)
    _clear_document_body(template_doc)

    samples = _cover_letter_samples(sample_doc, style_anchor.metadata.get("style_anchor_structure") or {})
    content = tiptap_json.get("content", []) if isinstance(tiptap_json, dict) else []

    for paragraph in samples["letterhead"]:
        _copy_sample_paragraph(template_doc, paragraph)
    template_doc.add_paragraph("")
    template_doc.add_paragraph("")

    state = {
        "body_started": False,
        "saw_re_line": False,
        "in_exhibit_section": False,
        "exhibit_counter": 1,
        "closing_requested": False,
        "signature_started": False,
        "intro_paragraphs_written": 0,
    }

    for node in content:
        _process_cover_letter_node(template_doc, node, samples, state)

    template_doc.add_paragraph("")
    _copy_sample_paragraph(template_doc, samples["closing_salutation"])

    template_doc.add_paragraph("")
    template_doc.add_paragraph("")
    _render_signature_block(template_doc, samples["signature"])

    buffer = io.BytesIO()
    template_doc.save(buffer)
    buffer.seek(0)
    return buffer


def _clear_document_body(doc):
    body = doc._element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def _cover_letter_samples(doc, structure):
    items = []
    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").replace("\xa0", " ").strip()
        if text:
            items.append({"paragraph": paragraph, "text": text})

    def find_text(prefix, *, from_end=False):
        iterable = reversed(items) if from_end else items
        for item in iterable:
            if item["text"].startswith(prefix):
                return item["paragraph"]
        return None

    letterhead_lines = structure.get("letterhead_lines") or []
    signature_lines = structure.get("signature_lines") or []

    letterhead = []
    for line in letterhead_lines:
        paragraph = find_text(line)
        if paragraph:
            letterhead.append(paragraph)
    if not letterhead:
        letterhead = [item["paragraph"] for item in items[:6]]

    signature = []
    for line in signature_lines:
        paragraph = find_text(line, from_end=True)
        if paragraph:
            signature.append(paragraph)
    if not signature:
        signature = [item["paragraph"] for item in items[-6:]]

    service_line = find_text("Via ") or items[min(6, len(items) - 1)]["paragraph"]
    date_line = next(
        (
            item["paragraph"]
            for item in items
            if re.match(r"^([A-Z][a-z]+ \d{1,2}, \d{4}|\[Date\]|\d{1,2}/\d{1,2}/\d{2,4})", item["text"])
        ),
        service_line,
    )
    re_line = find_text("RE:") or items[min(8, len(items) - 1)]["paragraph"]
    subject_line = find_text("Applicant:") or items[min(9, len(items) - 1)]["paragraph"]
    salutation = find_text("Dear ") or items[min(10, len(items) - 1)]["paragraph"]
    intro = next(
        (
            item["paragraph"]
            for item in items
            if item["text"].startswith("Undersigned counsel")
        ),
        items[min(11, len(items) - 1)]["paragraph"],
    )
    body = next(
        (
            item["paragraph"]
            for item in items
            if item["paragraph"].paragraph_format.first_line_indent
            and int(item["paragraph"].paragraph_format.first_line_indent) > 0
        ),
        intro,
    )
    section_heading = next(
        (
            item["paragraph"]
            for item in items
            if item["paragraph"].style.name == "List Paragraph"
            and any(run.bold for run in item["paragraph"].runs)
        ),
        body,
    )
    exhibit_category = next(
        (
            item["paragraph"]
            for item in items
            if item["text"].endswith(":")
            and any(run.underline and run.italic for run in item["paragraph"].runs)
        ),
        section_heading,
    )
    closing_salutation = find_text("Respectfully submitted") or items[-7]["paragraph"]

    return {
        "letterhead": letterhead,
        "signature": signature,
        "service_line": service_line,
        "date_line": date_line,
        "address_line": items[min(7, len(items) - 1)]["paragraph"],
        "re_line": re_line,
        "subject_line": subject_line,
        "salutation": salutation,
        "intro": intro,
        "body": body,
        "section_heading": section_heading,
        "exhibit_category": exhibit_category,
        "closing_salutation": closing_salutation,
    }


def _copy_paragraph_style(target, sample):
    target.style = sample.style
    target.alignment = sample.alignment
    source_format = sample.paragraph_format
    target_format = target.paragraph_format
    target_format.left_indent = source_format.left_indent
    target_format.right_indent = source_format.right_indent
    target_format.first_line_indent = source_format.first_line_indent
    target_format.space_before = source_format.space_before
    target_format.space_after = source_format.space_after
    target_format.line_spacing = source_format.line_spacing
    target_format.line_spacing_rule = source_format.line_spacing_rule
    target_format.keep_together = source_format.keep_together
    target_format.keep_with_next = source_format.keep_with_next
    target_format.page_break_before = source_format.page_break_before
    target_format.widow_control = source_format.widow_control


def _copy_run_style(target_run, sample_run):
    if not sample_run:
        return
    target_run.bold = sample_run.bold
    target_run.italic = sample_run.italic
    target_run.underline = sample_run.underline
    target_run.style = sample_run.style
    target_run.font.name = sample_run.font.name
    target_run.font.size = sample_run.font.size
    target_run.font.bold = sample_run.font.bold
    target_run.font.italic = sample_run.font.italic
    target_run.font.underline = sample_run.font.underline
    target_run.font.all_caps = sample_run.font.all_caps
    target_run.font.small_caps = sample_run.font.small_caps
    target_run.font.superscript = sample_run.font.superscript
    target_run.font.subscript = sample_run.font.subscript
    target_run.font.strike = sample_run.font.strike


def _copy_sample_paragraph(doc, sample):
    paragraph = doc.add_paragraph()
    _copy_paragraph_style(paragraph, sample)
    for sample_run in sample.runs:
        run = paragraph.add_run(sample_run.text)
        _copy_run_style(run, sample_run)
    return paragraph


def _render_signature_block(doc, signature_samples):
    for sample in signature_samples:
        line = (sample.text or "").replace("\xa0", " ").strip()
        if not line:
            continue
        paragraph = doc.add_paragraph()
        _copy_paragraph_style(paragraph, sample)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(line)
        _copy_run_style(run, sample.runs[0] if sample.runs else None)


def _apply_text_alignment(paragraph, alignment):
    if alignment == "center":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "right":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == "justify":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif alignment == "left":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _apply_document_page_setup(doc, document_metadata):
    page_setup = (document_metadata or {}).get("page_setup") or {}
    if not page_setup:
        return
    for section in doc.sections:
        if page_setup.get("page_width_pt") is not None:
            section.page_width = Pt(page_setup["page_width_pt"])
        if page_setup.get("page_height_pt") is not None:
            section.page_height = Pt(page_setup["page_height_pt"])
        if page_setup.get("left_margin_pt") is not None:
            section.left_margin = Pt(page_setup["left_margin_pt"])
        if page_setup.get("right_margin_pt") is not None:
            section.right_margin = Pt(page_setup["right_margin_pt"])
        if page_setup.get("top_margin_pt") is not None:
            section.top_margin = Pt(page_setup["top_margin_pt"])
        if page_setup.get("bottom_margin_pt") is not None:
            section.bottom_margin = Pt(page_setup["bottom_margin_pt"])
        if page_setup.get("header_distance_pt") is not None:
            section.header_distance = Pt(page_setup["header_distance_pt"])
        if page_setup.get("footer_distance_pt") is not None:
            section.footer_distance = Pt(page_setup["footer_distance_pt"])


def _resolve_style_name(doc, attrs, fallback=""):
    word_style = attrs.get("word_style") or {}
    candidate_names = [
        word_style.get("name") or "",
        word_style.get("style_id") or "",
        fallback or "",
    ]
    for name in candidate_names:
        if not name:
            continue
        try:
            doc.styles[name]
            return name
        except KeyError:
            continue
    return fallback or ""


def _apply_paragraph_metrics(paragraph, metrics):
    if not metrics:
        return
    fmt = paragraph.paragraph_format
    if metrics.get("left_indent_pt") is not None:
        fmt.left_indent = Pt(metrics["left_indent_pt"])
    if metrics.get("right_indent_pt") is not None:
        fmt.right_indent = Pt(metrics["right_indent_pt"])
    if metrics.get("first_line_indent_pt") is not None:
        fmt.first_line_indent = Pt(metrics["first_line_indent_pt"])
    if metrics.get("space_before_pt") is not None:
        fmt.space_before = Pt(metrics["space_before_pt"])
    if metrics.get("space_after_pt") is not None:
        fmt.space_after = Pt(metrics["space_after_pt"])
    if metrics.get("line_spacing") is not None:
        fmt.line_spacing = metrics["line_spacing"]
    line_spacing_rule = str(metrics.get("line_spacing_rule") or "").upper()
    if "DOUBLE" in line_spacing_rule:
        fmt.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    elif "SINGLE" in line_spacing_rule:
        fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif "ONE_POINT_FIVE" in line_spacing_rule:
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    elif "EXACTLY" in line_spacing_rule:
        fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    elif "AT_LEAST" in line_spacing_rule:
        fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    if metrics.get("keep_together") is not None:
        fmt.keep_together = metrics["keep_together"]
    if metrics.get("keep_with_next") is not None:
        fmt.keep_with_next = metrics["keep_with_next"]
    if metrics.get("page_break_before") is not None:
        fmt.page_break_before = metrics["page_break_before"]
    if metrics.get("widow_control") is not None:
        fmt.widow_control = metrics["widow_control"]
    _apply_tab_stops(fmt, metrics.get("tab_stops") or [])


def _apply_tab_stops(paragraph_format, tab_stops):
    if not tab_stops:
        return
    alignment_map = {
        "left": WD_TAB_ALIGNMENT.LEFT,
        "center": WD_TAB_ALIGNMENT.CENTER,
        "right": WD_TAB_ALIGNMENT.RIGHT,
        "decimal": WD_TAB_ALIGNMENT.DECIMAL,
        "bar": WD_TAB_ALIGNMENT.BAR,
    }
    leader_map = {
        "spaces": WD_TAB_LEADER.SPACES,
        "dots": WD_TAB_LEADER.DOTS,
        "dashes": WD_TAB_LEADER.DASHES,
        "lines": WD_TAB_LEADER.LINES,
        "heavy": WD_TAB_LEADER.HEAVY,
        "middle_dot": WD_TAB_LEADER.MIDDLE_DOT,
    }
    for stop in tab_stops:
        position = stop.get("position_pt")
        if position is None:
            continue
        paragraph_format.tab_stops.add_tab_stop(
            Pt(position),
            alignment=alignment_map.get(stop.get("alignment"), WD_TAB_ALIGNMENT.LEFT),
            leader=leader_map.get(stop.get("leader"), WD_TAB_LEADER.SPACES),
        )


def _apply_paragraph_attrs(paragraph, doc, node, fallback_style="", default_indent=False):
    attrs = node.get("attrs", {}) or {}
    style_name = _resolve_style_name(doc, attrs, fallback=fallback_style)
    if style_name:
        paragraph.style = style_name
    _apply_paragraph_metrics(paragraph, attrs.get("paragraph_metrics") or {})
    _apply_text_alignment(paragraph, attrs.get("textAlign") or (attrs.get("paragraph_metrics") or {}).get("alignment"))
    if default_indent and not attrs.get("paragraph_metrics"):
        paragraph.paragraph_format.left_indent = Inches(0.5)


def _apply_run_metrics(run, metrics):
    if not metrics:
        return
    if metrics.get("font_name"):
        run.font.name = metrics["font_name"]
    if metrics.get("font_size_pt") is not None:
        run.font.size = Pt(metrics["font_size_pt"])
    if metrics.get("bold") is not None:
        run.font.bold = metrics["bold"]
        run.bold = metrics["bold"]
    if metrics.get("italic") is not None:
        run.font.italic = metrics["italic"]
        run.italic = metrics["italic"]
    if metrics.get("underline") is not None:
        run.font.underline = metrics["underline"]
        run.underline = metrics["underline"]
    if metrics.get("all_caps") is not None:
        run.font.all_caps = metrics["all_caps"]
    if metrics.get("small_caps") is not None:
        run.font.small_caps = metrics["small_caps"]
    if metrics.get("strike") is not None:
        run.font.strike = metrics["strike"]
    if metrics.get("superscript") is not None:
        run.font.superscript = metrics["superscript"]
    if metrics.get("subscript") is not None:
        run.font.subscript = metrics["subscript"]


def _apply_inline_parts_with_sample(paragraph, inline_parts, sample_run):
    for part in inline_parts:
        part_type = part.get("type")
        if part_type == "text":
            run = paragraph.add_run(part.get("text", ""))
            _copy_run_style(run, sample_run)
            marks = part.get("marks", {})
            _apply_run_metrics(run, (marks.get("wordRun") or {}).get("run_metrics") or {})
            if "bold" in marks:
                run.bold = True
            if "italic" in marks:
                run.italic = True
            if "underline" in marks:
                run.underline = True
            if "strike" in marks:
                run.font.strike = True
            if "superscript" in marks:
                run.font.superscript = True
            if "subscript" in marks:
                run.font.subscript = True
        elif part_type == "hardBreak":
            paragraph.add_run().add_break(WD_BREAK.LINE)


def _add_paragraph_from_parts(doc, sample, inline_parts, alignment=None):
    paragraph = doc.add_paragraph()
    _copy_paragraph_style(paragraph, sample)
    if alignment:
        _apply_text_alignment(paragraph, alignment)
    sample_run = sample.runs[0] if sample.runs else None
    _apply_inline_parts_with_sample(paragraph, inline_parts, sample_run)
    return paragraph


def _node_plain_text(node):
    text_parts = []
    for part in _extract_inline_parts(node):
        if part.get("type") == "text":
            text_parts.append(part.get("text", ""))
    return "".join(text_parts).replace("\xa0", " ").strip()


def _looks_like_subject_detail(text):
    prefixes = (
        "Applicant:",
        "Petitioner:",
        "Beneficiary:",
        "Joint Petitioner:",
        "Receipt Number:",
        "A#:",
        "A#:",
        "Form Type:",
        "Qualifying Relative:",
    )
    return text.startswith(prefixes)


def _looks_like_signature_detail(text):
    lowered = text.lower()
    return lowered.startswith("christopher hammond") or lowered.startswith("chris hammond law firm") or lowered.startswith("attorney for ") or lowered.startswith("24 greenway plaza") or lowered.startswith("houston, texas") or lowered.startswith("tel.") or lowered.startswith("email.")


def _is_exhibit_heading(text):
    lowered = text.lower()
    return any(
        phrase in lowered
        for phrase in [
            "enclosed documentation",
            "evidence submitted",
            "supporting evidence",
            "exhibits",
            "evidence index",
            "please find enclosed",
            "please find attached",
        ]
    )


def _is_exhibit_category(text):
    return text.endswith(":") and len(text) <= 120


def _render_re_line(doc, sample, text, alignment=None):
    paragraph = doc.add_paragraph()
    _copy_paragraph_style(paragraph, sample)
    if alignment:
        _apply_text_alignment(paragraph, alignment)
    subject = text[3:].strip() if text.startswith("RE:") else text.strip()
    label_run = paragraph.add_run("RE:")
    _copy_run_style(label_run, sample.runs[0] if sample.runs else None)
    tab_run = paragraph.add_run("\t")
    _copy_run_style(tab_run, sample.runs[1] if len(sample.runs) > 1 else sample.runs[0] if sample.runs else None)
    subject_run = paragraph.add_run(subject)
    subject_sample = sample.runs[2] if len(sample.runs) > 2 else sample.runs[-1] if sample.runs else None
    _copy_run_style(subject_run, subject_sample)
    return paragraph


def _render_subject_detail(doc, sample, text, alignment=None):
    paragraph = doc.add_paragraph()
    _copy_paragraph_style(paragraph, sample)
    if alignment:
        _apply_text_alignment(paragraph, alignment)
    run = paragraph.add_run(text)
    _copy_run_style(run, sample.runs[0] if sample.runs else None)
    return paragraph


def _render_exhibit_heading(doc, sample, text, alignment=None):
    paragraph = doc.add_paragraph()
    _copy_paragraph_style(paragraph, sample)
    if alignment:
        _apply_text_alignment(paragraph, alignment)
    run = paragraph.add_run(text if text.endswith(":") else f"{text}:")
    _copy_run_style(run, sample.runs[0] if sample.runs else None)
    return paragraph


def _remove_table_borders(table):
    table_element = table._tbl
    table_properties = table_element.tblPr
    borders = table_properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_element = borders.find(qn(f"w:{edge}"))
        if edge_element is None:
            edge_element = OxmlElement(f"w:{edge}")
            borders.append(edge_element)
        edge_element.set(qn("w:val"), "nil")


def _render_exhibit_table(doc, items, samples, counter_start):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _remove_table_borders(table)
    label_sample = samples["section_heading"].runs[0] if samples["section_heading"].runs else None
    text_sample = samples["body"].runs[0] if samples["body"].runs else None
    for offset, item in enumerate(items):
        row = table.add_row()
        row.cells[0].width = Inches(1.55)
        row.cells[1].width = Inches(4.95)
        label_paragraph = row.cells[0].paragraphs[0]
        label_paragraph.style = doc.styles["Normal"]
        label_run = label_paragraph.add_run(f"EXHIBIT {counter_start + offset}")
        _copy_run_style(label_run, label_sample)
        label_run.bold = True

        text_paragraph = row.cells[1].paragraphs[0]
        _copy_paragraph_style(text_paragraph, samples["body"])
        text_paragraph.paragraph_format.first_line_indent = Inches(0)
        text_paragraph.paragraph_format.left_indent = Inches(0)
        text_run = text_paragraph.add_run(item.rstrip(";") + ";")
        _copy_run_style(text_run, text_sample)
    return counter_start + len(items)


def _process_cover_letter_node(doc, node, samples, state):
    node_type = node.get("type", "")
    text = _node_plain_text(node)
    alignment = node.get("attrs", {}).get("textAlign")

    if state["signature_started"]:
        return

    if node_type == "paragraph":
        if not text:
            doc.add_paragraph("")
            return
        if _looks_like_signature_detail(text):
            state["signature_started"] = True
            return
        if text.startswith("Respectfully submitted"):
            state["closing_requested"] = True
            return

        inline_parts = _extract_inline_parts(node)
        if not state["body_started"]:
            if text.startswith("Via "):
                _add_paragraph_from_parts(doc, samples["service_line"], inline_parts, alignment=alignment)
                return
            if DATE_PATTERN.match(text):
                _add_paragraph_from_parts(doc, samples["date_line"], inline_parts, alignment=alignment)
                return
            if text.startswith("RE:"):
                _render_re_line(doc, samples["re_line"], text, alignment=alignment)
                state["saw_re_line"] = True
                return
            if _looks_like_subject_detail(text):
                _render_subject_detail(doc, samples["subject_line"], text, alignment=alignment)
                return
            if text.lower().startswith("dear "):
                _add_paragraph_from_parts(doc, samples["salutation"], inline_parts, alignment=alignment)
                state["body_started"] = True
                return
            _add_paragraph_from_parts(doc, samples["address_line"], inline_parts, alignment=alignment)
            return

        if _is_exhibit_heading(text):
            state["in_exhibit_section"] = True
            _add_paragraph_from_parts(doc, samples["body"], inline_parts, alignment=alignment)
            return
        if state["in_exhibit_section"] and _is_exhibit_category(text):
            _render_exhibit_heading(doc, samples["exhibit_category"], text, alignment=alignment)
            return

        sample = samples["intro"] if state["intro_paragraphs_written"] < 2 else samples["body"]
        _add_paragraph_from_parts(doc, sample, inline_parts, alignment=alignment)
        state["intro_paragraphs_written"] += 1
        return

    if node_type == "heading":
        if not text:
            return
        level = node.get("attrs", {}).get("level", 1)
        if not state["body_started"] and level == 1 and not state["saw_re_line"]:
            _render_re_line(doc, samples["re_line"], f"RE: {text.title() if text.isupper() else text}", alignment=alignment)
            state["saw_re_line"] = True
            return
        if _is_exhibit_heading(text):
            state["in_exhibit_section"] = True
        if state["in_exhibit_section"] and level >= 3:
            _render_exhibit_heading(doc, samples["exhibit_category"], text, alignment=alignment)
            return
        paragraph = doc.add_paragraph()
        _copy_paragraph_style(paragraph, samples["section_heading"])
        if alignment:
            _apply_text_alignment(paragraph, alignment)
        run = paragraph.add_run(text)
        _copy_run_style(run, samples["section_heading"].runs[0] if samples["section_heading"].runs else None)
        run.bold = True
        state["body_started"] = True
        state["intro_paragraphs_written"] += 1
        return

    if node_type == "bulletList":
        items = []
        for item in node.get("content", []):
            for child in item.get("content", []):
                if child.get("type") == "paragraph":
                    item_text = _node_plain_text(child)
                    if item_text:
                        items.append(item_text)
        if not items:
            return
        if state["in_exhibit_section"]:
            state["exhibit_counter"] = _render_exhibit_table(doc, items, samples, state["exhibit_counter"])
            return
        for item_text in items:
            paragraph = doc.add_paragraph(style="List Bullet")
            run = paragraph.add_run(item_text)
            _copy_run_style(run, samples["body"].runs[0] if samples["body"].runs else None)
        return

    if node_type == "orderedList":
        items = []
        for item in node.get("content", []):
            for child in item.get("content", []):
                if child.get("type") == "paragraph":
                    item_text = _node_plain_text(child)
                    if item_text:
                        items.append(item_text)
        for number, item_text in enumerate(items, start=1):
            paragraph = doc.add_paragraph(style="List Number")
            run = paragraph.add_run(item_text)
            _copy_run_style(run, samples["body"].runs[0] if samples["body"].runs else None)
        return

    if node_type == "table":
        _process_table(doc, node, FORMAT_PRESETS["cover_letter"], [])
        return

    if node_type == "pageBreak":
        paragraph = doc.add_paragraph()
        paragraph.add_run().add_break(WD_BREAK.PAGE)
        return


def _process_node(doc, node, preset, para_counter, footnotes, *, preserve_template_styles=False):
    """Process a single Tiptap node into docx elements."""
    node_type = node.get("type", "")

    if node_type == "heading":
        level = node.get("attrs", {}).get("level", 1)
        level = min(max(level, 1), 3)
        inline_parts = _extract_inline_parts(node)
        fallback_style = f"Heading {level}"
        p = doc.add_paragraph(style=fallback_style if fallback_style in doc.styles else None)
        _apply_paragraph_attrs(p, doc, node, fallback_style=fallback_style)
        _apply_text_parts(p, inline_parts, footnotes)

    elif node_type == "paragraph":
        inline_parts = _extract_inline_parts(node)
        # Skip empty paragraphs (just add spacing)
        if not inline_parts or all(
            part.get("type") == "text" and part.get("text", "").strip() == ""
            for part in inline_parts
        ):
            doc.add_paragraph("")
            return

        if preset.get("numbered_paragraphs"):
            para_counter[0] += 1
            p = doc.add_paragraph()
            _apply_paragraph_attrs(p, doc, node, fallback_style="Normal")
            run = p.add_run(f"{para_counter[0]}. ")
            run.font.name = preset["font_name"]
            run.font.size = Pt(preset["font_size"])
            _apply_text_parts(p, inline_parts, footnotes)
        else:
            p = doc.add_paragraph()
            _apply_paragraph_attrs(p, doc, node, fallback_style="Normal")
            _apply_text_parts(p, inline_parts, footnotes)

    elif node_type == "bulletList":
        for item in node.get("content", []):
            _process_list_item(
                doc,
                item,
                node,
                preset,
                footnotes,
                bullet=True,
                preserve_template_styles=preserve_template_styles,
            )

    elif node_type == "orderedList":
        for i, item in enumerate(node.get("content", []), 1):
            _process_list_item(
                doc,
                item,
                node,
                preset,
                footnotes,
                bullet=False,
                number=i,
                preserve_template_styles=preserve_template_styles,
            )

    elif node_type == "blockquote":
        for child in node.get("content", []):
            if child.get("type") == "paragraph":
                inline_parts = _extract_inline_parts(child)
                p = doc.add_paragraph()
                _apply_paragraph_attrs(p, doc, child, fallback_style="Normal", default_indent=True)
                _apply_text_parts(p, inline_parts, footnotes)

    elif node_type == "horizontalRule":
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run("_" * 60)
        run.font.size = Pt(8)
        run.font.color.rgb = None

    elif node_type == "pageBreak":
        p = doc.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)

    elif node_type == "table":
        _process_table(doc, node, preset, footnotes, preserve_template_styles=preserve_template_styles)

    elif node_type == "hardBreak":
        pass  # Handled within text extraction


def _process_list_item(doc, item, list_node, preset, footnotes, bullet=True, number=None, preserve_template_styles=False):
    """Process a list item node."""
    list_attrs = (list_node.get("attrs") or {}) if isinstance(list_node, dict) else {}
    fallback_style = (
        ((list_attrs.get("list_identity") or {}).get("style_name") or "")
        or ("List Bullet" if bullet else "List Number")
    )
    try:
        doc.styles[fallback_style]
    except KeyError:
        fallback_style = "List Bullet" if bullet else "List Number"
    for child in item.get("content", []):
        if child.get("type") == "paragraph":
            inline_parts = _extract_inline_parts(child)
            p = doc.add_paragraph(style=fallback_style)
            _apply_text_parts(p, inline_parts, footnotes)
            _apply_paragraph_attrs(
                p,
                doc,
                child,
                fallback_style=fallback_style,
            )
            if not preserve_template_styles:
                p.style.font.name = preset["font_name"]
                p.style.font.size = Pt(preset["font_size"])


def _process_table(doc, node, preset, footnotes, preserve_template_styles=False):
    """Process a table node."""
    rows_data = node.get("content", [])
    if not rows_data:
        return

    # Determine dimensions
    num_rows = len(rows_data)
    num_cols = max(
        len(row.get("content", [])) for row in rows_data
    ) if rows_data else 1

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table_style = ((node.get("attrs") or {}).get("word_style") or {}).get("name") or "Table Grid"
    try:
        table.style = table_style
    except KeyError:
        if not preserve_template_styles:
            table.style = "Table Grid"
    if not preserve_template_styles and table.style is None:
        table.style = "Table Grid"
    table_alignment = ((node.get("attrs") or {}).get("paragraph_metrics") or {}).get("alignment")
    if table_alignment == "center":
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    elif table_alignment == "right":
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    else:
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for r_idx, row_node in enumerate(rows_data):
        cells = row_node.get("content", [])
        for c_idx, cell_node in enumerate(cells):
            if c_idx < num_cols:
                cell = table.rows[r_idx].cells[c_idx]
                # Clear default paragraph
                cell.paragraphs[0].clear()
                for child in cell_node.get("content", []):
                    if child.get("type") == "paragraph":
                        inline_parts = _extract_inline_parts(child)
                        p = cell.paragraphs[0] if not cell.paragraphs[0].text else cell.add_paragraph()
                        _apply_paragraph_attrs(p, doc, child, fallback_style="Normal")
                        _apply_text_parts(p, inline_parts, footnotes)


def _extract_inline_parts(node):
    """
    Extract inline content and formatting from a node.
    Returns list of dicts describing text, breaks, and footnote refs.
    """
    parts = []
    for child in node.get("content", []):
        child_type = child.get("type")
        if child_type == "text":
            text = child.get("text", "")
            marks = {}
            for mark in child.get("marks", []):
                mark_type = mark.get("type", "")
                marks[mark_type] = mark.get("attrs", {}) if mark.get("attrs") is not None else {}
            parts.append({"type": "text", "text": text, "marks": marks})
        elif child_type == "hardBreak":
            parts.append({"type": "hardBreak"})
        elif child_type == "footnoteReference":
            attrs = child.get("attrs", {})
            parts.append(
                {
                    "type": "footnoteReference",
                    "number": attrs.get("number"),
                    "text": attrs.get("text") or "",
                }
            )
    return parts


def _apply_text_parts(paragraph, inline_parts, footnotes):
    """Apply inline parts with formatting to a paragraph."""
    for part in inline_parts:
        part_type = part.get("type")
        if part_type == "text":
            run = paragraph.add_run(part.get("text", ""))
            marks = part.get("marks", {})
            _apply_run_metrics(run, (marks.get("wordRun") or {}).get("run_metrics") or {})
            if "bold" in marks:
                run.bold = True
            if "italic" in marks:
                run.italic = True
            if "underline" in marks:
                run.underline = True
            if "strike" in marks:
                run.font.strike = True
            if "superscript" in marks:
                run.font.superscript = True
            if "subscript" in marks:
                run.font.subscript = True
        elif part_type == "hardBreak":
            paragraph.add_run().add_break(WD_BREAK.LINE)
        elif part_type == "footnoteReference":
            number = _coerce_footnote_number(part.get("number"), len(footnotes) + 1)
            marker = paragraph.add_run(f"[{number}]")
            marker.font.superscript = True
            _register_footnote(footnotes, number, part.get("text") or "")


def _register_footnote(footnotes, number, text):
    number = _coerce_footnote_number(number, len(footnotes) + 1)
    for item in footnotes:
        if int(item["number"]) == int(number):
            if text and not item["text"]:
                item["text"] = text
            return
    footnotes.append({"number": number, "text": text})


def _coerce_footnote_number(number, fallback):
    try:
        return int(number)
    except (TypeError, ValueError):
        return int(fallback)


def _collect_footnotes(tiptap_json):
    footnotes = []
    nodes = tiptap_json.get("content", []) if isinstance(tiptap_json, dict) else []

    def walk(node):
        if isinstance(node, list):
            for item in node:
                walk(item)
            return
        if not isinstance(node, dict):
            return

        if node.get("type") == "footnoteReference":
            attrs = node.get("attrs", {})
            number = attrs.get("number")
            text = attrs.get("text") or ""
            if number is not None:
                _register_footnote(footnotes, number, text)
            return

        walk(node.get("content", []))

    walk(nodes)
    footnotes.sort(key=lambda x: x["number"])
    return footnotes


def _render_footnotes_html(footnotes):
    if not footnotes:
        return ""
    items = "".join(
        f"<li><span class=\"fn-marker\">[{item['number']}]</span> {escape(item['text'] or '')}</li>"
        for item in footnotes
    )
    return (
        "<section class=\"footnotes\">"
        "<h3>Footnotes</h3>"
        f"<ol>{items}</ol>"
        "</section>"
    )


def tiptap_to_pdf(tiptap_json, title="Document", export_format="court_brief"):
    """Convert Tiptap JSON content to PDF using WeasyPrint."""
    try:
        from weasyprint import HTML
    except ImportError as exc:
        raise RuntimeError(
            "PDF export requires WeasyPrint. Install with: pip install weasyprint"
        ) from exc

    html = tiptap_to_html(tiptap_json, title=title, export_format=export_format)
    buffer = io.BytesIO()
    HTML(string=html).write_pdf(target=buffer)
    buffer.seek(0)
    return buffer


def tiptap_to_html(tiptap_json, title="Document", export_format="court_brief"):
    preset = FORMAT_PRESETS.get(export_format, FORMAT_PRESETS["court_brief"])
    line_height = 2.0 if preset["line_spacing"] == WD_LINE_SPACING.DOUBLE else 1.2
    body_nodes = tiptap_json.get("content", []) if isinstance(tiptap_json, dict) else []
    body_html = "".join(_render_node_html(node) for node in body_nodes)
    footnotes_html = _render_footnotes_html(_collect_footnotes(tiptap_json))

    return f"""<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <title>{escape(title)}</title>
  <style>
    @page {{
      margin: {preset["margin_inches"]}in;
      size: letter;
    }}
    body {{
      font-family: "{preset["font_name"]}", serif;
      font-size: {preset["font_size"]}pt;
      line-height: {line_height};
      color: #1f2937;
    }}
    h1 {{ font-size: 14pt; margin: 12pt 0 6pt; }}
    h2 {{ font-size: 13pt; margin: 10pt 0 5pt; }}
    h3 {{ font-size: 12pt; margin: 8pt 0 4pt; }}
    p {{ margin: 4pt 0; }}
    blockquote {{
      border-left: 2px solid #9ca3af;
      margin: 8pt 0;
      padding-left: 10pt;
      color: #4b5563;
    }}
    hr {{ border: none; border-top: 1px solid #d1d5db; margin: 10pt 0; }}
    .page-break {{ break-after: page; page-break-after: always; border: 0; margin: 0; }}
    table {{ width: 100%; border-collapse: collapse; margin: 8pt 0; }}
    td, th {{ border: 1px solid #9ca3af; padding: 4pt; vertical-align: top; }}
    ol, ul {{ margin: 6pt 0 6pt 20pt; }}
    .fn-ref {{ font-size: 0.85em; }}
    .footnotes {{ margin-top: 20pt; border-top: 1px solid #d1d5db; padding-top: 8pt; }}
    .footnotes h3 {{ font-size: 11pt; margin: 0 0 6pt; }}
    .footnotes ol {{ margin-left: 18pt; }}
    .footnotes li {{ margin-bottom: 4pt; }}
    .fn-marker {{ font-size: 0.9em; }}
  </style>
</head>
<body>
{body_html}
{footnotes_html}
</body>
</html>"""


def _render_node_html(node):
    node_type = node.get("type")
    if node_type == "paragraph":
        attrs = node.get("attrs", {})
        align = attrs.get("textAlign")
        style = f' style="text-align:{align};"' if align else ""
        return f"<p{style}>{_render_inline_html(node.get('content', []))}</p>"
    if node_type == "heading":
        attrs = node.get("attrs", {})
        level = attrs.get("level", 1)
        level = max(1, min(3, int(level)))
        align = attrs.get("textAlign")
        style = f' style="text-align:{align};"' if align else ""
        return f"<h{level}{style}>{_render_inline_html(node.get('content', []))}</h{level}>"
    if node_type == "bulletList":
        items = "".join(_render_list_item_html(item) for item in node.get("content", []))
        return f"<ul>{items}</ul>"
    if node_type == "orderedList":
        items = "".join(_render_list_item_html(item) for item in node.get("content", []))
        return f"<ol>{items}</ol>"
    if node_type == "blockquote":
        inner = "".join(_render_node_html(child) for child in node.get("content", []))
        return f"<blockquote>{inner}</blockquote>"
    if node_type == "horizontalRule":
        return "<hr>"
    if node_type == "pageBreak":
        return '<hr class="page-break">'
    if node_type == "table":
        rows = []
        for row in node.get("content", []):
            cols = []
            for cell in row.get("content", []):
                tag = "th" if cell.get("type") == "tableHeader" else "td"
                children = "".join(_render_node_html(child) for child in cell.get("content", []))
                cols.append(f"<{tag}>{children}</{tag}>")
            rows.append(f"<tr>{''.join(cols)}</tr>")
        return f"<table>{''.join(rows)}</table>"
    return ""


def _render_list_item_html(item):
    children = item.get("content", [])
    parts = []
    for child in children:
        if child.get("type") == "paragraph":
            parts.append(_render_inline_html(child.get("content", [])))
        else:
            parts.append(_render_node_html(child))
    return f"<li>{''.join(parts)}</li>"


def _render_inline_html(content):
    pieces = []
    for node in content:
        node_type = node.get("type")
        if node_type == "text":
            text = escape(node.get("text", ""))
            for mark in node.get("marks", []):
                mark_type = mark.get("type")
                if mark_type == "bold":
                    text = f"<strong>{text}</strong>"
                elif mark_type == "italic":
                    text = f"<em>{text}</em>"
                elif mark_type == "underline":
                    text = f"<u>{text}</u>"
                elif mark_type == "strike":
                    text = f"<s>{text}</s>"
                elif mark_type == "superscript":
                    text = f"<sup>{text}</sup>"
                elif mark_type == "subscript":
                    text = f"<sub>{text}</sub>"
                elif mark_type == "link":
                    href = escape((mark.get("attrs") or {}).get("href") or "#")
                    text = f'<a href="{href}">{text}</a>'
            pieces.append(text)
        elif node_type == "footnoteReference":
            attrs = node.get("attrs", {})
            number = attrs.get("number") or "?"
            pieces.append(f"<sup class=\"fn-ref\">[{escape(str(number))}]</sup>")
        elif node_type == "hardBreak":
            pieces.append("<br>")
    return "".join(pieces)
