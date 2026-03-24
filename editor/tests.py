from unittest.mock import patch
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
import shutil
import tempfile

from django.contrib.auth.models import User
from django.core.management import call_command
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from .agent_service import (
    AGENT_FINALIZATION_MAX_OUTPUT_TOKENS,
    AGENT_FINALIZATION_REASONING_EFFORT,
    AgentConfigurationError,
    DocumentResearchAgent,
    _search_client_files_for_agent,
    _client_file_function_tools,
    _extract_json_object,
    _extract_output_text,
    _extract_hosted_tool_calls,
    _knowledge_function_tools,
    _normalize_edit_result,
    _normalize_mcp_server_url,
    _request_requirements_block,
    _requested_full_text_sources,
)
from .export import tiptap_to_docx, tiptap_to_html
from .import_service import import_docx_package, import_docx_to_tiptap
from .models import (
    Document,
    DocumentClientFile,
    DocumentResearchMessage,
    DocumentResearchRun,
    DocumentResearchSession,
    DocumentVersion,
    DocumentType,
    Exemplar,
    WritingWorkspace,
    WorkspaceResearchMessage,
    WorkspaceResearchRun,
    WorkspaceResearchSession,
)
from .openai_file_service import analyze_client_file_with_input_file, sync_client_file_openai_index
from .proof_service import ProofRenderError, SofficeRenderBackend, render_document_proof


def _sample_tiptap(text):
    return {
        "type": "doc",
        "content": [
            {
                "type": "paragraph",
                "content": [
                    {
                        "type": "text",
                        "text": text,
                    }
                ],
            }
        ],
    }


def _sample_cover_letter():
    return {
        "type": "doc",
        "content": [
            {"type": "paragraph", "content": [{"type": "text", "text": "Via FedEx"}]},
            {"type": "paragraph", "content": []},
            {"type": "paragraph", "content": [{"type": "text", "text": "March 7, 2026"}]},
            {"type": "paragraph", "content": []},
            {"type": "paragraph", "content": [{"type": "text", "text": "U.S. Citizenship and Immigration Services"}]},
            {"type": "paragraph", "content": [{"type": "text", "text": "Dallas Lockbox"}]},
            {"type": "paragraph", "content": [{"type": "text", "text": "P.O. Box 660867"}]},
            {"type": "paragraph", "content": [{"type": "text", "text": "Dallas, TX 75266"}]},
            {"type": "paragraph", "content": []},
            {"type": "paragraph", "content": [{"type": "text", "text": "RE: Form I-130, Petition for Alien Relative"}]},
            {"type": "paragraph", "content": [{"type": "text", "text": "Petitioner: Jane Doe"}]},
            {"type": "paragraph", "content": [{"type": "text", "text": "Beneficiary: John Doe"}]},
            {"type": "paragraph", "content": []},
            {"type": "paragraph", "content": [{"type": "text", "text": "Dear USCIS Officer:"}]},
            {
                "type": "paragraph",
                "content": [{"type": "text", "text": "Please find enclosed this Form I-130 filing and supporting documentation."}],
            },
            {"type": "heading", "attrs": {"level": 2}, "content": [{"type": "text", "text": "A. Case Summary"}]},
            {
                "type": "paragraph",
                "content": [{"type": "text", "text": "The petitioner is a U.S. citizen seeking classification for her spouse."}],
            },
            {
                "type": "paragraph",
                "content": [{"type": "text", "text": "Please find enclosed the following documents in support of this filing:"}],
            },
            {"type": "heading", "attrs": {"level": 3}, "content": [{"type": "text", "text": "Forms"}]},
            {
                "type": "bulletList",
                "content": [
                    {"type": "listItem", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "Form I-130, Petition for Alien Relative"}]}]},
                    {"type": "listItem", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "Form G-28, Notice of Entry of Appearance"}]}]},
                ],
            },
            {"type": "heading", "attrs": {"level": 3}, "content": [{"type": "text", "text": "Supporting Evidence"}]},
            {
                "type": "bulletList",
                "content": [
                    {"type": "listItem", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "Marriage certificate"}]}]},
                    {"type": "listItem", "content": [{"type": "paragraph", "content": [{"type": "text", "text": "Petitioner passport copy"}]}]},
                ],
            },
            {"type": "paragraph", "content": [{"type": "text", "text": "Respectfully submitted,"}]},
            {"type": "paragraph", "content": [{"type": "text", "text": "Chris Hammond"}]},
            {"type": "paragraph", "content": [{"type": "text", "text": "Attorney for Petitioner"}]},
        ],
    }


def _build_docx_bytes(*, margin_inches=1.0):
    doc = DocxDocument()
    section = doc.sections[0]
    section.left_margin = Inches(margin_inches)
    section.right_margin = Inches(margin_inches)
    doc.styles["Normal"].font.name = "Courier New"

    heading = doc.add_heading("Imported Brief Heading", level=1)
    heading.runs[0].bold = True

    paragraph = doc.add_paragraph()
    paragraph.add_run("This is a ")
    paragraph.add_run("bold").bold = True
    paragraph.add_run(" paragraph.")

    bullet = doc.add_paragraph(style="List Bullet")
    bullet.add_run("First exhibit item")

    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].paragraphs[0].add_run("Exhibit").bold = True
    table.rows[0].cells[1].paragraphs[0].add_run("Description").bold = True
    table.rows[1].cells[0].text = "A"
    table.rows[1].cells[1].text = "Marriage certificate"

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


class AgentResearchViewsTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="lawyer", password="secret")
        self.client.force_login(self.user)
        self.document_type = DocumentType.objects.create(
            name="Defensive Asylum Brief",
            slug="defensive-asylum-brief",
            category="brief",
            template_content=_sample_tiptap("Template"),
        )
        self.document = Document.objects.create(
            title="Test Brief",
            document_type=self.document_type,
            content=_sample_tiptap("The client fears return because gang threats escalated after reporting extortion."),
            created_by=self.user,
        )

    def test_agent_session_bootstraps_empty_session(self):
        response = self.client.get(
            reverse("research_agent_session", kwargs={"doc_id": self.document.id})
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["messages"], [])
        self.assertTrue(DocumentResearchSession.objects.filter(document=self.document, user=self.user).exists())

    @patch("editor.agent_views.DocumentResearchAgent")
    def test_agent_chat_starts_background_run_and_persists_user_message(self, agent_cls):
        agent = agent_cls.return_value

        def fake_start_chat_run(*, run, **kwargs):
            run.status = "in_progress"
            run.stage = "waiting_openai"
            run.response_id = "resp_test_123"
            run.response_count = 1
            run.save(update_fields=["status", "stage", "response_id", "response_count", "updated_at"])
            return run

        agent.start_chat_run.side_effect = fake_start_chat_run

        response = self.client.post(
            reverse("research_agent_chat", kwargs={"doc_id": self.document.id}),
            data={
                "message": "What precedent should I add here?",
                "selected_text": "The client fears return because gang threats escalated.",
            },
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 202)
        payload = response.json()
        self.assertEqual(payload["user_message"]["content"], "What precedent should I add here?")
        self.assertEqual(payload["run"]["status"], "in_progress")
        self.assertEqual(payload["run"]["response_id"], "resp_test_123")

        session = DocumentResearchSession.objects.get(document=self.document, user=self.user)
        self.assertEqual(session.last_response_id, "")
        self.assertEqual(session.messages.count(), 1)
        self.assertTrue(DocumentResearchRun.objects.filter(session=session, response_id="resp_test_123").exists())

    def test_agent_chat_returns_conflict_when_run_already_active(self):
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="chat",
            status="in_progress",
            stage="waiting_openai",
        )

        response = self.client.post(
            reverse("research_agent_chat", kwargs={"doc_id": self.document.id}),
            data={"message": "Can you help with this paragraph?"},
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 409)
        payload = response.json()
        self.assertEqual(payload["run"]["id"], str(run.public_id))

    def test_agent_reset_clears_messages(self):
        session = DocumentResearchSession.objects.create(
            document=self.document,
            user=self.user,
            last_response_id="resp_old",
        )
        DocumentResearchMessage.objects.create(session=session, role="user", content="Question")
        DocumentResearchMessage.objects.create(session=session, role="assistant", content="Answer", response_id="resp_old")

        response = self.client.post(
            reverse("research_agent_reset", kwargs={"doc_id": self.document.id})
        )

        self.assertEqual(response.status_code, 200)
        session.refresh_from_db()
        self.assertEqual(session.last_response_id, "")
        self.assertEqual(session.messages.count(), 0)

    @patch("editor.agent_views.DocumentResearchAgent")
    def test_agent_chat_returns_json_for_unexpected_exception(self, agent_cls):
        agent_cls.return_value.start_chat_run.side_effect = RuntimeError("boom")

        response = self.client.post(
            reverse("research_agent_chat", kwargs={"doc_id": self.document.id}),
            data={"message": "Help me with this paragraph."},
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 500)
        self.assertEqual(
            response.json()["error"],
            "The agent failed unexpectedly while starting.",
        )

    @patch("editor.agent_views.DocumentResearchAgent")
    def test_agent_chat_returns_json_for_constructor_configuration_error(self, agent_cls):
        agent_cls.side_effect = AgentConfigurationError("OPENAI_API_KEY is not configured.")

        response = self.client.post(
            reverse("research_agent_chat", kwargs={"doc_id": self.document.id}),
            data={"message": "Help me with this paragraph."},
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 503)
        self.assertEqual(
            response.json()["error"],
            "OPENAI_API_KEY is not configured.",
        )

    @patch("editor.agent_views.DocumentResearchAgent")
    def test_agent_run_status_persists_completed_chat_message(self, agent_cls):
        agent = agent_cls.return_value
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        user_message = DocumentResearchMessage.objects.create(
            session=session,
            role="user",
            content="Can you strengthen this section?",
            selection_text="The client reported gang extortion to police.",
        )
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="chat",
            status="in_progress",
            stage="waiting_openai",
            user_message=user_message,
        )

        def fake_advance_run(*, run):
            run.status = "completed"
            run.stage = "completed"
            run.response_id = "resp_chat_done"
            run.result_payload = {
                "answer": "Matter of C-T-L- supports the nexus rule here.",
                "response_id": "resp_chat_done",
                "tool_calls": [{"source": "biaedge", "type": "mcp_call", "name": "search_cases"}],
                "citations": [],
                "used_tools": ["biaedge"],
                "metadata": {"model": "gpt-5.4"},
            }
            run.save(update_fields=["status", "stage", "response_id", "result_payload", "updated_at"])
            return run

        agent.advance_run.side_effect = fake_advance_run

        response = self.client.get(
            reverse("research_agent_run", kwargs={"run_id": run.public_id})
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["run"]["status"], "completed")
        self.assertEqual(payload["assistant_message"]["content"], "Matter of C-T-L- supports the nexus rule here.")

        run.refresh_from_db()
        session.refresh_from_db()
        self.assertIsNotNone(run.assistant_message_id)
        self.assertEqual(session.last_response_id, "resp_chat_done")
        self.assertTrue(
            DocumentResearchMessage.objects.filter(session=session, role="assistant", response_id="resp_chat_done").exists()
        )

    @patch("editor.agent_views._persist_chat_completion", side_effect=RuntimeError("persist boom"))
    def test_agent_run_status_returns_fallback_chat_message_when_persist_fails(self, persist_completion):
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        user_message = DocumentResearchMessage.objects.create(
            session=session,
            role="user",
            content="Can you strengthen this section?",
            selection_text="The client reported gang extortion to police.",
        )
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="chat",
            status="completed",
            stage="completed",
            user_message=user_message,
            response_id="resp_chat_done",
            result_payload={
                "answer": "Matter of C-T-L- supports the nexus rule here.",
                "response_id": "resp_chat_done",
                "tool_calls": [{"source": "biaedge", "type": "mcp_call", "name": "search_cases"}],
                "citations": [],
                "used_tools": ["biaedge"],
                "metadata": {"model": "gpt-5.4"},
            },
        )

        response = self.client.get(
            reverse("research_agent_run", kwargs={"run_id": run.public_id})
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["run"]["status"], "completed")
        self.assertEqual(payload["assistant_message"]["content"], "Matter of C-T-L- supports the nexus rule here.")
        self.assertEqual(payload["assistant_message"]["response_id"], "resp_chat_done")
        persist_completion.assert_called_once()

        run.refresh_from_db()
        self.assertTrue(run.metadata.get("assistant_persist_failed"))
        self.assertIsNone(run.assistant_message_id)

    @patch("editor.agent_views.DocumentResearchAgent")
    def test_agent_suggest_returns_json_for_constructor_configuration_error(self, agent_cls):
        agent_cls.side_effect = AgentConfigurationError("OPENAI_API_KEY is not configured.")

        response = self.client.post(
            reverse("research_agent_suggest", kwargs={"doc_id": self.document.id}),
            data={"selected_text": "Selected text"},
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 503)
        self.assertEqual(
            response.json()["error"],
            "OPENAI_API_KEY is not configured.",
        )

    @patch("editor.agent_views.DocumentResearchAgent")
    def test_agent_suggest_starts_background_run(self, agent_cls):
        agent = agent_cls.return_value

        def fake_start_suggest_run(*, run, **kwargs):
            run.status = "in_progress"
            run.stage = "waiting_openai"
            run.response_id = "resp_suggest_1"
            run.response_count = 1
            run.save(update_fields=["status", "stage", "response_id", "response_count", "updated_at"])
            return run

        agent.start_suggest_run.side_effect = fake_start_suggest_run

        response = self.client.post(
            reverse("research_agent_suggest", kwargs={"doc_id": self.document.id}),
            data={
                "selected_text": "The threats were because he reported gang extortion.",
                "focus_note": "Find the best nexus authority.",
            },
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 202)
        payload = response.json()
        self.assertEqual(payload["run"]["status"], "in_progress")
        self.assertEqual(payload["run"]["response_id"], "resp_suggest_1")

    @patch("editor.agent_views.DocumentResearchAgent")
    def test_agent_run_status_returns_completed_suggest_payload(self, agent_cls):
        agent = agent_cls.return_value
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="suggest",
            status="in_progress",
            stage="waiting_openai",
        )

        def fake_advance_run(*, run):
            run.status = "completed"
            run.stage = "completed"
            run.response_id = "resp_suggest_done"
            run.result_payload = {
                "selection_summary": "Nexus support for gang-based persecution.",
                "draft_gap": "The paragraph needs controlling nexus authority and one central reason language.",
                "authorities": [
                    {
                        "kind": "case",
                        "title": "Matter of C-T-L-",
                        "citation": "25 I&N Dec. 341 (BIA 2010)",
                        "document_id": 101,
                        "reference_id": None,
                        "precedential_status": "precedential",
                        "validity_status": "good_law",
                        "relevance": "Explains the one central reason standard.",
                        "suggested_use": "Use it to frame the nexus paragraph.",
                        "pinpoint": "One protected ground must be one central reason for the harm.",
                    }
                ],
                "search_notes": "Searched precedential BIA nexus authorities.",
                "next_questions": ["Do you also need PSG-specific authority?"],
                "tool_calls": [{"source": "biaedge", "name": "search_cases", "type": "mcp_call"}],
                "citations": [],
                "raw_answer": "{}",
            }
            run.save(update_fields=["status", "stage", "response_id", "result_payload", "updated_at"])
            return run

        agent.advance_run.side_effect = fake_advance_run

        response = self.client.get(
            reverse("research_agent_run", kwargs={"run_id": run.public_id})
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["run"]["status"], "completed")
        self.assertEqual(payload["suggest_result"]["selection_summary"], "Nexus support for gang-based persecution.")
        self.assertEqual(payload["suggest_result"]["authorities"][0]["citation"], "25 I&N Dec. 341 (BIA 2010)")

    @patch("editor.agent_views.DocumentResearchAgent")
    def test_agent_edit_starts_background_run(self, agent_cls):
        agent = agent_cls.return_value

        def fake_start_edit_run(*, run, **kwargs):
            run.status = "in_progress"
            run.stage = "waiting_openai"
            run.response_id = "resp_edit_1"
            run.response_count = 1
            run.save(update_fields=["status", "stage", "response_id", "response_count", "updated_at"])
            return run

        agent.start_edit_run.side_effect = fake_start_edit_run

        response = self.client.post(
            reverse("research_agent_edit", kwargs={"doc_id": self.document.id}),
            data={
                "instruction": "Rewrite this paragraph to tighten the nexus analysis.",
                "selected_text": "The client fears return because gang threats escalated.",
                "selection_from": 5,
                "selection_to": 62,
            },
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 202)
        payload = response.json()
        self.assertEqual(payload["run"]["status"], "in_progress")
        self.assertEqual(payload["run"]["response_id"], "resp_edit_1")

    @patch("editor.agent_views.DocumentResearchAgent")
    def test_agent_run_status_returns_completed_edit_payload(self, agent_cls):
        agent = agent_cls.return_value
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="edit",
            status="in_progress",
            stage="waiting_openai",
        )

        def fake_advance_run(*, run):
            run.status = "completed"
            run.stage = "completed"
            run.response_id = "resp_edit_done"
            run.result_payload = {
                "edit_summary": "Tighten the nexus paragraph.",
                "rationale": "This version states the legal standard first and then ties it to the facts.",
                "operation": "replace_selection",
                "target_text": "The client fears return because gang threats escalated.",
                "proposed_text": "The record shows the gang threatened the client because he reported the extortion to police.",
                "notes": "Uses a clearer causal link.",
                "selected_text": "The client fears return because gang threats escalated.",
                "selection_from": 5,
                "selection_to": 62,
                "selection_required": True,
                "tool_calls": [{"source": "biaedge", "name": "search_cases", "type": "mcp_call"}],
                "citations": [],
                "raw_answer": "{}",
            }
            run.save(update_fields=["status", "stage", "response_id", "result_payload", "updated_at"])
            return run

        agent.advance_run.side_effect = fake_advance_run

        response = self.client.get(
            reverse("research_agent_run", kwargs={"run_id": run.public_id})
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["run"]["status"], "completed")
        self.assertEqual(payload["edit_result"]["operation"], "replace_selection")
        self.assertIn("legal standard first", payload["edit_result"]["rationale"])

    def test_agent_apply_edit_snapshots_current_content_and_saves_new_content(self):
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="edit",
            status="completed",
            stage="completed",
            result_payload={
                "edit_summary": "Strengthen nexus paragraph",
                "operation": "replace_selection",
                "proposed_text": "Updated paragraph text.",
            },
        )
        current_content = _sample_tiptap("Original paragraph text.")
        new_content = _sample_tiptap("Updated paragraph text.")

        response = self.client.post(
            reverse("research_agent_apply_edit", kwargs={"doc_id": self.document.id}),
            data={
                "run_id": str(run.public_id),
                "current_content": current_content,
                "new_content": new_content,
            },
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.document.refresh_from_db()
        run.refresh_from_db()
        self.assertEqual(self.document.content, new_content)
        self.assertEqual(run.result_payload["applied_at"], payload["edit_result"]["applied_at"])
        snapshot = DocumentVersion.objects.get(id=payload["version"]["id"])
        self.assertEqual(snapshot.content, current_content)
        self.assertTrue(snapshot.label.startswith("Before agent edit - Strengthen nexus paragraph"))

    def test_client_document_upload_and_list_for_current_document(self):
        upload = SimpleUploadedFile(
            "police-report.txt",
            b"On March 1, 2024, the client reported the threats to police.",
            content_type="text/plain",
        )

        upload_response = self.client.post(
            reverse("document_client_file_upload", kwargs={"doc_id": self.document.id}),
            data={"file": upload, "title": "Police Report"},
        )

        self.assertEqual(upload_response.status_code, 200)
        payload = upload_response.json()["client_file"]
        self.assertEqual(payload["title"], "Police Report")
        self.assertIn("reported the threats", payload["extracted_text"])

        list_response = self.client.get(
            reverse("document_client_file_list", kwargs={"doc_id": self.document.id}),
            data={"q": "police"},
        )

        self.assertEqual(list_response.status_code, 200)
        results = list_response.json()["results"]
        self.assertEqual(len(results), 1)
        self.assertEqual(results[0]["title"], "Police Report")

    @patch("editor.document_file_views.sync_client_file_openai_index")
    @patch("editor.document_file_views.extract_text_from_file")
    def test_client_document_upload_marks_scanned_pdf_as_openai_analysis_ready(self, extract_text, sync_index):
        extract_text.return_value = ""

        def fake_sync_index(client_file):
            return {
                **(client_file.metadata or {}),
                "char_count": 0,
                "text_extracted": False,
                "scan_candidate": True,
                "openai_index_status": "completed",
                "openai_file_id": "file_client_scan",
                "openai_vector_store_id": "vs_client_scan",
                "openai_vector_store_file_id": "vsfile_client_scan",
            }

        sync_index.side_effect = fake_sync_index
        upload = SimpleUploadedFile(
            "cbp-scan.pdf",
            b"%PDF-1.4 fake",
            content_type="application/pdf",
        )

        response = self.client.post(
            reverse("document_client_file_upload", kwargs={"doc_id": self.document.id}),
            data={"file": upload, "title": "CBP Scan"},
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()["client_file"]
        self.assertEqual(payload["metadata"]["openai_index_status"], "completed")
        self.assertTrue(payload["metadata"]["scan_candidate"])
        self.assertIn("OpenAI document analysis", payload["metadata"]["warning"])


class EditorViewTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="editor-user", password="secret")
        self.client.force_login(self.user)
        self.document_type = DocumentType.objects.create(
            name="I-130 Cover Letter",
            slug="i-130-cover-letter-editor",
            category="cover_letter",
            template_content=_sample_tiptap("Template"),
        )
        self.document = Document.objects.create(
            title="Draft Cover Letter",
            document_type=self.document_type,
            content=_sample_tiptap("Opening paragraph."),
            created_by=self.user,
        )

    def test_editor_page_renders_enhanced_writing_controls(self):
        response = self.client.get(reverse("editor", kwargs={"doc_id": self.document.id}))

        self.assertEqual(response.status_code, 200)
        self.assertContains(response, 'id="block-style-select"', html=False)
        self.assertContains(response, 'id="outline-sidebar"', html=False)
        self.assertContains(response, 'id="editor-statusbar"', html=False)
        self.assertContains(response, 'id="focus-toggle"', html=False)
        self.assertContains(response, '@tiptap/extension-text-align', html=False)
        self.assertContains(response, 'Cmd/Ctrl+K link', html=False)


class WordAddinViewsTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="word-addin-user", password="secret")
        self.client.force_login(self.user)
        self.document_type = DocumentType.objects.create(
            name="Asylum Brief",
            slug="asylum-brief",
            category="brief",
            template_content=_sample_tiptap("Template"),
            export_format="court_brief",
            order=1,
        )
        self.second_document_type = DocumentType.objects.create(
            name="Waiver Cover Letter",
            slug="waiver-cover-letter",
            category="cover_letter",
            template_content=_sample_tiptap("Template"),
            export_format="cover_letter",
            order=2,
        )

    def test_word_addin_manifest_uses_public_commands_page(self):
        response = self.client.get(reverse("word_addin_manifest"))

        self.assertEqual(response.status_code, 200)
        self.assertIn("text/xml", response["Content-Type"])
        body = response.content.decode("utf-8")
        self.assertIn("/word-addin/commands/", body)
        self.assertIn("/word-addin/taskpane/", body)
        self.assertIn("<AppDomain>http://testserver</AppDomain>", body)

    def test_word_addin_document_types_returns_expected_fields(self):
        response = self.client.get(reverse("word_addin_document_types"))

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(
            payload["document_types"],
            [
                {
                    "id": self.document_type.id,
                    "name": "Asylum Brief",
                    "slug": "asylum-brief",
                    "category": "brief",
                    "export_format": "court_brief",
                },
                {
                    "id": self.second_document_type.id,
                    "name": "Waiver Cover Letter",
                    "slug": "waiver-cover-letter",
                    "category": "cover_letter",
                    "export_format": "cover_letter",
                },
            ],
        )

    def test_word_addin_workspace_bootstrap_creates_workspace_and_session(self):
        response = self.client.post(
            reverse("word_addin_workspace_bootstrap"),
            data={
                "document_title": "Cancellation Motion Draft",
                "document_type_slug": self.document_type.slug,
                "external_document_key": "word://draft-123",
                "persistent": False,
            },
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        workspace = WritingWorkspace.objects.get(id=payload["workspace"]["id"])
        session = WorkspaceResearchSession.objects.get(workspace=workspace, user=self.user)

        self.assertEqual(workspace.title, "Cancellation Motion Draft")
        self.assertEqual(workspace.document_type, self.document_type)
        self.assertEqual(workspace.external_document_key, "word://draft-123")
        self.assertFalse(workspace.metadata["persistent"])
        self.assertEqual(payload["session"]["id"], session.id)
        self.assertFalse(payload["persistent"])

    def test_word_addin_workspace_bootstrap_updates_existing_workspace(self):
        workspace = WritingWorkspace.objects.create(
            user=self.user,
            kind="word_addin",
            title="Old Title",
            document_type=self.document_type,
            external_document_key="word://old",
            metadata={"persistent": True},
        )

        response = self.client.post(
            reverse("word_addin_workspace_bootstrap"),
            data={
                "workspace_id": str(workspace.id),
                "document_title": "Updated Title",
                "document_type_slug": self.second_document_type.slug,
                "external_document_key": "word://updated",
                "persistent": True,
            },
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        workspace.refresh_from_db()
        self.assertEqual(workspace.title, "Updated Title")
        self.assertEqual(workspace.document_type, self.second_document_type)
        self.assertEqual(workspace.external_document_key, "word://updated")
        self.assertTrue(WorkspaceResearchSession.objects.filter(workspace=workspace, user=self.user).exists())

    def test_word_addin_workspace_session_returns_messages_and_latest_runs(self):
        workspace = WritingWorkspace.objects.create(
            user=self.user,
            kind="word_addin",
            title="Word Draft",
            document_type=self.document_type,
        )
        session = WorkspaceResearchSession.objects.create(workspace=workspace, user=self.user)
        WorkspaceResearchMessage.objects.create(
            session=session,
            role="user",
            content="What authority do I need here?",
            selection_text="The gang threatened the client after he reported extortion.",
        )
        assistant_message = WorkspaceResearchMessage.objects.create(
            session=session,
            role="assistant",
            content="Matter of C-T-L- is the starting point.",
            citations=[{"citation": "25 I&N Dec. 341 (BIA 2010)"}],
        )
        chat_run = WorkspaceResearchRun.objects.create(
            session=session,
            mode="chat",
            status="completed",
            result_payload={"answer": "Matter of C-T-L- is the starting point."},
            assistant_message=assistant_message,
        )
        suggest_run = WorkspaceResearchRun.objects.create(
            session=session,
            mode="suggest",
            status="completed",
            result_payload={
                "authorities": [
                    {
                        "citation": "25 I&N Dec. 341 (BIA 2010)",
                        "title": "Matter of C-T-L-",
                    }
                ]
            },
        )

        response = self.client.get(
            reverse("word_addin_workspace_session", kwargs={"workspace_id": workspace.id})
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["workspace"]["id"], str(workspace.id))
        self.assertEqual(len(payload["messages"]), 2)
        self.assertEqual(payload["latest_chat_run"]["id"], str(chat_run.public_id))
        self.assertEqual(payload["latest_suggest_run"]["id"], str(suggest_run.public_id))

    def test_word_addin_chat_persist_creates_messages_and_run(self):
        workspace = WritingWorkspace.objects.create(
            user=self.user,
            kind="word_addin",
            title="Word Draft",
            document_type=self.document_type,
        )

        response = self.client.post(
            reverse("word_addin_chat_persist", kwargs={"workspace_id": workspace.id}),
            data={
                "user_message": "What should I cite for nexus?",
                "assistant_message": "Matter of C-T-L- supports the one central reason analysis.",
                "selected_text": "The gang threatened him after he reported extortion.",
                "citations": [
                    {
                        "kind": "case",
                        "title": "Matter of C-T-L-",
                        "citation": "25 I&N Dec. 341 (BIA 2010)",
                        "document_id": 101,
                    }
                ],
                "bridge_job_id": "job-123",
                "metadata": {"source": "codex_bridge", "model": "gpt-5.4"},
                "status": "completed",
            },
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        run = WorkspaceResearchRun.objects.get(public_id=payload["run"]["id"])

        self.assertEqual(run.status, "completed")
        self.assertEqual(run.bridge_job_id, "job-123")
        self.assertEqual(run.user_message.content, "What should I cite for nexus?")
        self.assertEqual(run.assistant_message.content, "Matter of C-T-L- supports the one central reason analysis.")
        self.assertEqual(run.result_payload["citations"][0]["citation"], "25 I&N Dec. 341 (BIA 2010)")

    def test_word_addin_chat_persist_records_failed_run_without_assistant_message(self):
        workspace = WritingWorkspace.objects.create(
            user=self.user,
            kind="word_addin",
            title="Word Draft",
            document_type=self.document_type,
        )

        response = self.client.post(
            reverse("word_addin_chat_persist", kwargs={"workspace_id": workspace.id}),
            data={
                "user_message": "What should I cite for nexus?",
                "selected_text": "The gang threatened him after he reported extortion.",
                "bridge_job_id": "job-124",
                "status": "failed",
                "error_message": "Codex bridge failed.",
            },
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 502)
        run = WorkspaceResearchRun.objects.get(public_id=response.json()["run"]["id"])
        self.assertEqual(run.status, "failed")
        self.assertEqual(run.error_message, "Codex bridge failed.")
        self.assertIsNone(run.assistant_message)

    def test_word_addin_run_detail_uses_failed_status_code(self):
        workspace = WritingWorkspace.objects.create(
            user=self.user,
            kind="word_addin",
            title="Word Draft",
            document_type=self.document_type,
        )
        session = WorkspaceResearchSession.objects.create(workspace=workspace, user=self.user)
        run = WorkspaceResearchRun.objects.create(
            session=session,
            mode="suggest",
            status="failed",
            error_message="Codex bridge failed.",
        )

        response = self.client.get(
            reverse("word_addin_run_detail", kwargs={"run_id": run.public_id})
        )

        self.assertEqual(response.status_code, 502)
        self.assertEqual(response.json()["run"]["error_message"], "Codex bridge failed.")

    def test_word_addin_suggest_persist_requires_selected_text(self):
        workspace = WritingWorkspace.objects.create(
            user=self.user,
            kind="word_addin",
            title="Word Draft",
            document_type=self.document_type,
        )

        response = self.client.post(
            reverse("word_addin_suggest_persist", kwargs={"workspace_id": workspace.id}),
            data={"focus_note": "Nexus"},
            content_type="application/json",
        )

        self.assertEqual(response.status_code, 400)
        self.assertEqual(response.json()["error"], "selected_text is required.")

    def test_word_addin_citation_format_supports_styles(self):
        quote_response = self.client.post(
            reverse("word_addin_citation_format"),
            data={
                "style": "quote",
                "authority": {
                    "citation": "25 I&N Dec. 341 (BIA 2010)",
                    "pinpoint": "One protected ground must be one central reason for the harm.",
                },
            },
            content_type="application/json",
        )
        parenthetical_response = self.client.post(
            reverse("word_addin_citation_format"),
            data={
                "style": "parenthetical",
                "authority": {
                    "citation": "8 C.F.R. § 1208.13",
                    "suggested_use": "Sets out the asylum eligibility framework.",
                },
            },
            content_type="application/json",
        )

        self.assertEqual(quote_response.status_code, 200)
        self.assertEqual(
            quote_response.json()["plain_text"],
            "“One protected ground must be one central reason for the harm.” 25 I&N Dec. 341 (BIA 2010)",
        )
        self.assertEqual(
            parenthetical_response.json()["plain_text"],
            "8 C.F.R. § 1208.13 (Sets out the asylum eligibility framework.)",
        )


class ExportFormattingTests(TestCase):
    def test_docx_export_preserves_heading_alignment(self):
        content = {
            "type": "doc",
            "content": [
                {
                    "type": "heading",
                    "attrs": {"level": 2, "textAlign": "center"},
                    "content": [{"type": "text", "text": "Centered Heading"}],
                }
            ],
        }

        docx_buffer = tiptap_to_docx(content, title="Aligned Heading", export_format="court_brief")
        exported = DocxDocument(BytesIO(docx_buffer.getvalue()))

        self.assertEqual(exported.paragraphs[0].text, "Centered Heading")
        self.assertEqual(exported.paragraphs[0].alignment, WD_ALIGN_PARAGRAPH.CENTER)

    def test_html_export_preserves_heading_alignment(self):
        content = {
            "type": "doc",
            "content": [
                {
                    "type": "heading",
                    "attrs": {"level": 2, "textAlign": "center"},
                    "content": [{"type": "text", "text": "Centered Heading"}],
                }
            ],
        }

        html = tiptap_to_html(content, title="Aligned Heading", export_format="court_brief")
        self.assertIn('<h2 style="text-align:center;">Centered Heading</h2>', html)


class AgentServiceTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="agent_service_user", password="secret")
        self.document_type = DocumentType.objects.create(
            name="Agent Service Brief",
            slug="agent-service-brief",
            category="brief",
            template_content=_sample_tiptap("Template"),
        )
        self.document = Document.objects.create(
            title="Agent Service Test Document",
            document_type=self.document_type,
            content=_sample_tiptap("The client reported gang extortion to police."),
            created_by=self.user,
        )

    def test_knowledge_function_tools_use_valid_strict_required_lists(self):
        tools = {
            tool["name"]: tool
            for tool in _knowledge_function_tools()
        }

        search_schema = tools["search_exemplars"]["parameters"]
        self.assertEqual(
            search_schema["required"],
            ["query", "limit", "document_type_slug"],
        )
        self.assertEqual(
            sorted(search_schema["properties"].keys()),
            sorted(search_schema["required"]),
        )

        client_doc_schema = {
            tool["name"]: tool
            for tool in _client_file_function_tools()
        }["search_client_documents"]["parameters"]
        self.assertEqual(client_doc_schema["required"], ["query", "limit"])
        self.assertEqual(
            sorted(client_doc_schema["properties"].keys()),
            sorted(client_doc_schema["required"]),
        )
        analyze_schema = {
            tool["name"]: tool
            for tool in _client_file_function_tools()
        }["analyze_client_document"]["parameters"]
        self.assertEqual(analyze_schema["required"], ["file_id", "query"])
        self.assertEqual(
            sorted(analyze_schema["properties"].keys()),
            sorted(analyze_schema["required"]),
        )

    @patch("editor.agent_service._new_openai_client")
    def test_build_tools_omits_knowledge_functions_without_active_exemplars(self, new_client):
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )

        tools = agent._build_tools(mode="chat", include_mcp=False)

        self.assertFalse(any(tool.get("type") == "function" for tool in tools))

    @patch("editor.agent_service._new_openai_client")
    def test_build_tools_uses_low_context_web_search_for_suggest(self, new_client):
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )

        tools = agent._build_tools(mode="suggest", include_mcp=False)
        web_tool = next(tool for tool in tools if tool.get("type") == "web_search_preview")

        self.assertEqual(web_tool["search_context_size"], "low")

    @patch("editor.agent_service._new_openai_client")
    def test_build_tools_includes_client_document_functions_when_current_document_has_uploads(self, new_client):
        DocumentClientFile.objects.create(
            document=self.document,
            title="Police Report",
            original_file=SimpleUploadedFile("police-report.txt", b"Client reported the extortion to police."),
            extracted_text="Client reported the extortion to police.",
            uploaded_by=self.user,
            metadata={"filename": "police-report.txt", "extension": ".txt"},
        )
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )

        tools = agent._build_tools(mode="chat", include_mcp=False)
        function_names = sorted(tool.get("name") for tool in tools if tool.get("type") == "function")

        self.assertIn("analyze_client_document", function_names)
        self.assertIn("search_client_documents", function_names)
        self.assertIn("get_client_document", function_names)

    @patch("editor.agent_service.search_indexed_client_files")
    def test_search_client_documents_prefers_indexed_openai_results_when_available(self, search_indexed):
        client_file = DocumentClientFile.objects.create(
            document=self.document,
            title="CBP Inspection Record",
            original_file=SimpleUploadedFile("inspection.pdf", b"PDF bytes"),
            extracted_text="",
            uploaded_by=self.user,
            metadata={
                "filename": "inspection.pdf",
                "extension": ".pdf",
                "text_extracted": False,
                "scan_candidate": True,
                "openai_index_status": "completed",
            },
        )
        search_indexed.return_value = [
            {
                "id": client_file.id,
                "title": client_file.title,
                "filename": "inspection.pdf",
                "extension": ".pdf",
                "snippet": "The I-94 reflects admission on March 1, 2024.",
                "score": 0.9132,
                "text_extracted": False,
                "scan_candidate": True,
                "openai_index_status": "completed",
                "retrieval_source": "openai_vector_store",
            }
        ]

        result = _search_client_files_for_agent(
            document=self.document,
            query="inspection admission",
            limit=5,
        )

        self.assertEqual(result["results"][0]["id"], client_file.id)
        self.assertEqual(result["results"][0]["retrieval_source"], "openai_vector_store")
        search_indexed.assert_called_once_with(
            document=self.document,
            query="inspection admission",
            limit=5,
        )

    @patch("editor.agent_service.analyze_client_file_with_input_file")
    @patch("editor.agent_service._new_openai_client")
    def test_call_local_tool_can_analyze_original_client_document(self, new_client, analyze_client_file):
        client_file = DocumentClientFile.objects.create(
            document=self.document,
            title="Visa Stamp",
            original_file=SimpleUploadedFile("visa-stamp.pdf", b"PDF bytes"),
            extracted_text="",
            uploaded_by=self.user,
            metadata={"filename": "visa-stamp.pdf", "extension": ".pdf"},
        )
        analyze_client_file.return_value = {
            "id": client_file.id,
            "analysis": "The visa foil shows an F-1 classification issued in 2021.",
        }
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )

        result = agent._call_local_tool(
            name="analyze_client_document",
            arguments={"file_id": client_file.id, "query": "What does the visa foil show?"},
        )

        self.assertEqual(result["analysis"], "The visa foil shows an F-1 classification issued in 2021.")
        analyze_client_file.assert_called_once()

    @patch("editor.agent_service._new_openai_client")
    def test_local_function_budget_forces_final_response_instead_of_failing(self, new_client):
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="edit",
            status="in_progress",
            stage="running_tools",
            response_id="resp_local_tools",
            metadata=agent._initial_run_metadata(mode="edit", previous_response_id=""),
        )
        response = SimpleNamespace(id="resp_local_tools")
        function_call = SimpleNamespace(
            name="search_client_documents",
            call_id="call_client_doc_search",
            arguments='{"query":"draft statement","limit":2}',
        )
        queued_response = SimpleNamespace(id="resp_forced_final", status="queued", error=None, usage=None, output=[])

        with patch("editor.agent_service.AGENT_MAX_LOCAL_FUNCTION_ROUNDS", 0):
            with patch.object(agent, "_call_local_tool", return_value={"results": [{"id": 1, "title": "Draft Statement"}]}):
                with patch.object(agent, "_create_background_response", return_value=queued_response) as create_background:
                    updated = agent._continue_after_function_calls(
                        run=run,
                        response=response,
                        function_calls=[function_call],
                    )

        self.assertEqual(updated.status, "queued")
        self.assertEqual(updated.stage, "forcing_final")
        self.assertEqual(updated.response_id, "resp_forced_final")
        self.assertTrue(updated.metadata.get("allow_over_budget_finalization"))
        request = create_background.call_args.kwargs
        self.assertEqual(request["tools"], [])
        self.assertEqual(request["tool_choice"], "none")
        self.assertEqual(request["max_output_tokens"], AGENT_FINALIZATION_MAX_OUTPUT_TOKENS)
        self.assertEqual(request["reasoning_effort"], AGENT_FINALIZATION_REASONING_EFFORT)

    @patch("editor.agent_service._new_openai_client")
    def test_budget_error_allows_forced_finalization_after_local_tool_cap(self, new_client):
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="edit",
            status="in_progress",
            stage="forcing_final",
            local_function_rounds=9,
            metadata={"allow_over_budget_finalization": True},
        )

        with patch("editor.agent_service.AGENT_MAX_LOCAL_FUNCTION_ROUNDS", 8):
            self.assertEqual(agent._budget_error(run), "")

    @patch("editor.agent_service._new_openai_client")
    def test_budget_error_does_not_use_response_or_reasoning_caps(self, new_client):
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="chat",
            status="in_progress",
            stage="waiting_openai",
            response_count=99,
            usage={"reasoning_tokens": 999999, "total_tokens": 100},
            metadata=agent._initial_run_metadata(mode="chat", previous_response_id=""),
        )

        self.assertEqual(agent._budget_error(run), "")

    @patch("editor.agent_service._new_openai_client")
    def test_total_token_budget_queues_compact_finalization_instead_of_failing(self, new_client):
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="edit",
            status="in_progress",
            stage="running_tools",
            response_id="resp_budgeted",
            usage={"total_tokens": 130},
            tool_calls=[
                {
                    "source": "client_docs",
                    "type": "function_call",
                    "name": "get_client_document",
                    "status": "completed",
                    "arguments": {"file_id": 1},
                    "output_excerpt": '{"text":"Client states the officer requested the waiver at the interview."}',
                }
            ],
            metadata=agent._initial_run_metadata(mode="edit", previous_response_id=""),
        )
        queued_response = SimpleNamespace(id="resp_final_budget", status="queued", error=None, usage=None, output=[])

        with patch("editor.agent_service.AGENT_MAX_TOTAL_TOKENS", 120):
            with patch.object(agent, "_create_background_response", return_value=queued_response) as create_background:
                updated = agent.advance_run(run=run)

        self.assertEqual(updated.status, "queued")
        self.assertEqual(updated.stage, "forcing_final")
        self.assertEqual(updated.response_id, "resp_final_budget")
        self.assertTrue(updated.metadata.get("budget_recovery_attempted"))
        self.assertTrue(updated.metadata.get("allow_over_budget_finalization"))
        request = create_background.call_args.kwargs
        self.assertEqual(request["tools"], [])
        self.assertEqual(request["tool_choice"], "none")
        self.assertIn("verified research already gathered", request["instructions"].lower())

    @patch("editor.agent_service._new_openai_client")
    def test_run_response_loop_forces_final_answer_after_tool_only_round(self, new_client):
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )

        first_response = SimpleNamespace(
            id="resp_tool_only",
            status="completed",
            output_text="",
            output=[
                SimpleNamespace(
                    type="function_call",
                    name="search_exemplars",
                    status="completed",
                    call_id="call_123",
                    arguments='{"query":"nexus brief","limit":3,"document_type_slug":""}',
                )
            ],
        )
        final_response = SimpleNamespace(
            id="resp_final",
            status="completed",
            output_text="Matter of C-T-L- supports the nexus rule.",
            output=[
                SimpleNamespace(
                    type="message",
                    content=[
                        SimpleNamespace(
                            type="output_text",
                            text="Matter of C-T-L- supports the nexus rule.",
                            annotations=[],
                        )
                    ],
                )
            ],
        )

        with patch("editor.agent_service._MAX_FUNCTION_ROUNDS", 1):
            with patch.object(agent, "_create_response", side_effect=[first_response, final_response]) as create_response:
                result = agent._run_response_loop(
                    instructions="Test instructions",
                    input_payload="Test input",
                    tools=[{"type": "function", "name": "search_exemplars"}],
                    previous_response_id=None,
                    initial_tool_choice="auto",
                )

        self.assertEqual(result["answer"], "Matter of C-T-L- supports the nexus rule.")
        self.assertEqual(result["response_id"], "resp_final")
        self.assertEqual(create_response.call_count, 2)
        forced_request = create_response.call_args_list[1].kwargs
        self.assertEqual(forced_request["tools"], [])
        self.assertEqual(forced_request["previous_response_id"], "resp_tool_only")
        self.assertEqual(forced_request["tool_choice"], "none")
        self.assertEqual(forced_request["max_output_tokens"], AGENT_FINALIZATION_MAX_OUTPUT_TOKENS)
        self.assertEqual(forced_request["reasoning_effort"], AGENT_FINALIZATION_REASONING_EFFORT)

    def test_extract_output_text_reads_refusal_parts(self):
        response = SimpleNamespace(
            output_text="",
            output=[
                SimpleNamespace(
                    type="message",
                    content=[
                        SimpleNamespace(
                            type="refusal",
                            refusal="I can’t comply with that request.",
                        )
                    ],
                )
            ],
        )

        self.assertEqual(_extract_output_text(response), "I can’t comply with that request.")

    def test_extract_hosted_tool_calls_captures_mcp_output_excerpt(self):
        response = SimpleNamespace(
            output=[
                SimpleNamespace(
                    type="mcp_call",
                    name="get_reference",
                    status="completed",
                    arguments='{"reference_id": 324, "max_chars": 12000}',
                    output='{"title":"USCIS Policy Manual","text":"Generally, a CPR who initially files a waiver under one waiver filing basis may change to or add another waiver filing basis by making the request in writing."}',
                    error="",
                )
            ]
        )

        tool_calls = _extract_hosted_tool_calls(response)

        self.assertEqual(tool_calls[0]["name"], "get_reference")
        self.assertIn("output_excerpt", tool_calls[0])
        self.assertIn("waiver filing basis", tool_calls[0]["output_excerpt"])

    def test_requested_full_text_sources_detects_policy_and_case_quote_requests(self):
        sources = _requested_full_text_sources(
            "Find quotes from case law and quotes from the USCIS Policy Manual for this issue."
        )

        self.assertEqual(sources, {"case", "policy"})

    def test_request_requirements_block_guides_full_text_tools(self):
        note = _request_requirements_block(
            "Find quotes from case law and quotes from the USCIS Policy Manual that I can use."
        )

        self.assertIn("Turn requirements JSON:", note)
        self.assertIn('"requires_exact_quotes": true', note)
        self.assertIn('"required_full_text_sources": [', note)
        self.assertIn('"case"', note)
        self.assertIn('"policy"', note)
        self.assertIn("policy=search_references(source_code='uscis_pm')->get_reference", note)
        self.assertIn("case_law=get_document_text", note)

    @patch("editor.agent_service._new_openai_client")
    def test_start_edit_run_stores_structured_requirements(self, new_client):
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="edit",
            status="queued",
            stage="queued",
        )
        queued_response = SimpleNamespace(id="resp_edit_start", status="queued", error=None, usage=None, output=[])

        with patch.object(agent, "_build_tools", return_value=[]):
            with patch.object(agent, "_create_background_response", return_value=queued_response):
                updated = agent.start_edit_run(
                    run=run,
                    instruction="Find quotes from the USCIS Policy Manual and revise this section.",
                    selected_text="Selected paragraph text.",
                )

        self.assertEqual(updated.metadata.get("phase"), "research")
        requirements = updated.metadata.get("requirements") or {}
        self.assertEqual(requirements.get("mode"), "edit")
        self.assertEqual(requirements.get("output_format"), "edit_json")
        self.assertTrue(requirements.get("has_selected_text"))
        self.assertTrue(requirements.get("requires_exact_quotes"))
        self.assertIn("policy", requirements.get("required_full_text_sources") or [])

    @patch("editor.agent_service._new_openai_client")
    def test_record_response_artifacts_updates_evidence_pack(self, new_client):
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="chat",
            status="in_progress",
            stage="waiting_openai",
            metadata=agent._initial_run_metadata(mode="chat", previous_response_id=""),
        )
        response = SimpleNamespace(
            id="resp_evidence",
            usage=None,
            output=[
                SimpleNamespace(
                    type="mcp_call",
                    name="get_reference",
                    status="completed",
                    arguments='{"reference_id":324,"source_code":"uscis_pm"}',
                    output='{"title":"USCIS Policy Manual","text":"A CPR may change to or add another waiver filing basis by making the request in writing."}',
                    error="",
                )
            ],
        )

        agent._record_response_artifacts(run, response)

        evidence_pack = (run.metadata or {}).get("evidence_pack") or {}
        self.assertEqual(evidence_pack["counts"]["legal_authorities"], 1)
        self.assertEqual(evidence_pack["legal_authorities"][0]["tool"], "get_reference")
        self.assertIn("waiver filing basis", evidence_pack["legal_authorities"][0]["excerpt"])
        metrics = (run.metadata or {}).get("metrics") or {}
        self.assertEqual(metrics["tool_source_counts"]["biaedge"], 1)
        self.assertIn("biaedge", metrics["used_sources"])

    @patch("editor.agent_service._new_openai_client")
    def test_start_chat_run_includes_transcript_even_with_previous_response_id(self, new_client):
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="chat",
            status="queued",
            stage="queued",
        )
        transcript_messages = [
            SimpleNamespace(role="user", content="First question"),
            SimpleNamespace(role="assistant", content="First answer"),
        ]
        queued_response = SimpleNamespace(id="resp_chat_start", status="queued", error=None, usage=None, output=[])

        with patch.object(agent, "_build_tools", return_value=[]):
            with patch.object(agent, "_create_background_response", return_value=queued_response) as create_background:
                agent.start_chat_run(
                    run=run,
                    message="Follow-up question",
                    selected_text="",
                    previous_response_id="resp_prev",
                    transcript_messages=transcript_messages,
                )

        request = create_background.call_args.kwargs
        self.assertEqual(request["previous_response_id"], "resp_prev")
        self.assertIn("Conversation so far:", request["input_payload"])
        self.assertIn("First question", request["input_payload"])

    def test_normalize_edit_result_falls_back_to_append_without_selected_text_or_target(self):
        normalized = _normalize_edit_result(
            {
                "edit_summary": "Add a conclusion",
                "operation": "replace_selection",
                "proposed_text": "For these reasons, USCIS should approve the petition.",
            },
            request_payload={"selected_text": ""},
        )

        self.assertEqual(normalized["operation"], "append_to_document")

    def test_normalize_edit_result_keeps_targeted_replace_without_selected_text(self):
        normalized = _normalize_edit_result(
            {
                "edit_summary": "Revise the introduction paragraph",
                "operation": "replace_selection",
                "target_text": "Original introduction paragraph.",
                "proposed_text": "Updated introduction paragraph.",
            },
            request_payload={"selected_text": ""},
        )

        self.assertEqual(normalized["operation"], "replace_selection")
        self.assertEqual(normalized["target_text"], "Original introduction paragraph.")

    def test_normalize_edit_result_unescapes_double_escaped_newlines(self):
        normalized = _normalize_edit_result(
            {
                "edit_summary": "Draft sections A and B",
                "operation": "insert_after_selection",
                "target_text": "B. Eligibility and Procedural Posture",
                "proposed_text": "A. Case Summary\\nThis filing submits Form I-485.\\n\\nB. Eligibility and Procedural Posture\\nThe applicant seeks adjustment of status.",
            },
            request_payload={"selected_text": ""},
        )

        self.assertEqual(
            normalized["proposed_text"],
            "A. Case Summary\nThis filing submits Form I-485.\n\nB. Eligibility and Procedural Posture\nThe applicant seeks adjustment of status.",
        )

    def test_extract_json_object_allows_control_characters_inside_strings(self):
        raw = (
            '{"edit_summary":"Replace the argument section.",'
            '"rationale":"Improve substance.",'
            '"operation":"replace_selection",'
            '"target_text":"II. Argument\tFirst, the documentary evidence...",'
            '"proposed_text":"Revised argument text.",'
            '"notes":""}'
        )

        parsed = _extract_json_object(raw)

        self.assertIsInstance(parsed, dict)
        self.assertEqual(parsed["operation"], "replace_selection")
        self.assertIn("II. Argument", parsed["target_text"])

    @patch("editor.agent_service._new_openai_client")
    def test_queue_json_repair_salvages_plain_text_edit_response(self, new_client):
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="edit",
            status="in_progress",
            stage="repairing_json",
            request_payload={
                "instruction": "Strengthen this paragraph.",
                "selected_text": "Original paragraph text.",
            },
            metadata=agent._initial_run_metadata(mode="edit", previous_response_id=""),
        )
        run.metadata["json_repair_attempted"] = True
        run.save(update_fields=["metadata", "updated_at"])

        response = SimpleNamespace(
            id="resp_edit_plain",
            output_text="Proposed text:\nRevised paragraph text.",
        )

        updated = agent._queue_json_repair(run=run, response=response)

        self.assertEqual(updated.status, "completed")
        self.assertEqual(updated.result_payload["operation"], "replace_selection")
        self.assertEqual(updated.result_payload["proposed_text"], "Revised paragraph text.")

    @patch("editor.agent_service._new_openai_client")
    def test_queue_json_repair_salvages_structural_edit_from_plain_text_response(self, new_client):
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="edit",
            status="in_progress",
            stage="repairing_json",
            request_payload={
                "instruction": "Add a paragraph before III. CONCLUSION.",
                "selected_text": "",
            },
            metadata=agent._initial_run_metadata(mode="edit", previous_response_id=""),
        )
        run.metadata["json_repair_attempted"] = True
        run.save(update_fields=["metadata", "updated_at"])

        response = SimpleNamespace(
            id="resp_edit_structural",
            output_text="Target:\nIII. CONCLUSION\n\nProposed text:\nFor these reasons, the petition should be granted.",
        )

        updated = agent._queue_json_repair(run=run, response=response)

        self.assertEqual(updated.status, "completed")
        self.assertEqual(updated.result_payload["operation"], "insert_before_selection")
        self.assertEqual(updated.result_payload["target_text"], "III. CONCLUSION")

    def test_normalize_edit_result_coerces_append_to_replace_when_selection_exists(self):
        normalized = _normalize_edit_result(
            {
                "edit_summary": "Revise the selected paragraph",
                "operation": "append_to_document",
                "proposed_text": "Revised paragraph text.",
            },
            request_payload={"selected_text": "Original selected paragraph."},
        )

        self.assertEqual(normalized["operation"], "replace_selection")

    @patch("editor.agent_service._new_openai_client")
    def test_start_edit_run_ignores_micro_selection_for_structural_request(self, new_client):
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="edit",
            status="queued",
            stage="queued",
        )
        queued_response = SimpleNamespace(id="resp_edit_start", status="queued", error=None, usage=None, output=[])

        with patch.object(agent, "_build_tools", return_value=[]):
            with patch.object(agent, "_create_background_response", return_value=queued_response) as create_background:
                updated = agent.start_edit_run(
                    run=run,
                    instruction="Look at the uploaded documents and draft section A and B of this cover letter.",
                    selected_text="e",
                    selection_from=210,
                    selection_to=211,
                )

        self.assertEqual(updated.request_payload["selected_text"], "")
        self.assertIsNone(updated.request_payload["selection_from"])
        self.assertIsNone(updated.request_payload["selection_to"])
        self.assertTrue(updated.metadata.get("selection_ignored"))
        request = create_background.call_args.kwargs
        self.assertIn("No text is currently selected.", request["input_payload"])

    @patch("editor.agent_service._new_openai_client")
    def test_start_edit_run_keeps_micro_selection_for_word_level_request(self, new_client):
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="edit",
            status="queued",
            stage="queued",
        )
        queued_response = SimpleNamespace(id="resp_edit_word", status="queued", error=None, usage=None, output=[])

        with patch.object(agent, "_build_tools", return_value=[]):
            with patch.object(agent, "_create_background_response", return_value=queued_response):
                updated = agent.start_edit_run(
                    run=run,
                    instruction="Fix the spelling of this word.",
                    selected_text="teh",
                    selection_from=42,
                    selection_to=45,
                )

        self.assertEqual(updated.request_payload["selected_text"], "teh")
        self.assertEqual(updated.request_payload["selection_from"], 42)
        self.assertEqual(updated.request_payload["selection_to"], 45)

    def test_normalize_mcp_server_url_adds_scheme_and_default_path(self):
        self.assertEqual(
            _normalize_mcp_server_url("biaedge-mcp.onrender.com"),
            "https://biaedge-mcp.onrender.com/mcp",
        )
        self.assertEqual(
            _normalize_mcp_server_url("https://biaedge-mcp.onrender.com"),
            "https://biaedge-mcp.onrender.com/mcp",
        )
        self.assertEqual(
            _normalize_mcp_server_url("http://localhost:8001"),
            "http://localhost:8001/mcp",
        )

    @patch("editor.agent_service._new_openai_client")
    def test_failed_status_with_tool_budget_attempts_final_recovery(self, new_client):
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="chat",
            status="in_progress",
            stage="waiting_openai",
            response_id="resp_failed",
            tool_calls=[
                {"source": "biaedge", "name": f"search_cases_{index}", "type": "mcp_call"}
                for index in range(24)
            ],
        )
        failed_response = SimpleNamespace(
            id="resp_failed",
            status="failed",
            error=None,
            incomplete_details=None,
            usage=None,
            output=[],
        )
        recovery_response = SimpleNamespace(
            id="resp_recover",
            status="queued",
            error=None,
            usage=None,
            output=[],
        )

        agent.client.responses.retrieve = lambda *args, **kwargs: failed_response
        with patch.object(agent, "_create_background_response", return_value=recovery_response) as create_background:
            updated = agent.advance_run(run=run)

        self.assertEqual(updated.status, "queued")
        self.assertEqual(updated.stage, "recovering_failure")
        self.assertEqual(updated.response_id, "resp_recover")
        recovery_request = create_background.call_args.kwargs
        self.assertEqual(recovery_request["tool_choice"], "none")
        self.assertEqual(recovery_request["max_output_tokens"], AGENT_FINALIZATION_MAX_OUTPUT_TOKENS)
        self.assertEqual(recovery_request["reasoning_effort"], AGENT_FINALIZATION_REASONING_EFFORT)

    @patch("editor.agent_service._new_openai_client")
    def test_queue_force_final_response_uses_compacted_tool_digest_without_previous_chain(self, new_client):
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="chat",
            status="in_progress",
            stage="waiting_openai",
            response_id="resp_current",
            request_payload={
                "message": "Find quotes for this paragraph.",
                "selected_text": "Selected paragraph text.",
            },
            tool_calls=[
                {
                    "source": "biaedge",
                    "type": "mcp_call",
                    "name": "get_reference",
                    "status": "completed",
                    "arguments": {"reference_id": 324},
                    "output_excerpt": '{"title":"USCIS Policy Manual","text":"A CPR may change to or add another waiver filing basis by making the request in writing."}',
                }
            ],
        )
        current_response = SimpleNamespace(id="resp_current")
        queued_response = SimpleNamespace(id="resp_final", status="queued", error=None, usage=None, output=[])

        with patch.object(agent, "_create_background_response", return_value=queued_response) as create_background:
            updated = agent._queue_force_final_response(run=run, response=current_response)

        self.assertEqual(updated.response_id, "resp_final")
        request = create_background.call_args.kwargs
        self.assertIsNone(request["previous_response_id"])
        self.assertEqual(request["tool_choice"], "none")
        self.assertIn("Evidence pack:", request["input_payload"])
        self.assertIn("waiver filing basis", request["input_payload"])
        self.assertEqual(updated.metadata["finalization_source"], "empty_response")
        self.assertEqual(updated.metadata["metrics"]["finalization_source"], "empty_response")

    @patch("editor.agent_service._new_openai_client")
    def test_finalize_chat_run_updates_session_last_response_id(self, new_client):
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="chat",
            status="in_progress",
            stage="waiting_openai",
            metadata=agent._initial_run_metadata(mode="chat", previous_response_id=""),
        )
        response = SimpleNamespace(id="resp_chat_final")

        completed = agent._finalize_chat_run(run=run, response=response, answer="Final answer text.")

        self.assertEqual(completed.result_payload["response_id"], "resp_chat_final")
        session.refresh_from_db()
        self.assertEqual(session.last_response_id, "resp_chat_final")

    @patch("editor.agent_service._new_openai_client")
    def test_failed_status_with_tool_budget_has_clearer_error_message(self, new_client):
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="chat",
            status="in_progress",
            stage="waiting_openai",
            response_id="resp_failed",
            tool_calls=[
                {"source": "biaedge", "name": f"search_cases_{index}", "type": "mcp_call"}
                for index in range(24)
            ],
            metadata={"failed_recovery_attempted": True},
        )
        failed_response = SimpleNamespace(
            id="resp_failed",
            status="failed",
            error=None,
            incomplete_details=None,
            usage=None,
            output=[],
        )

        agent.client.responses.retrieve = lambda *args, **kwargs: failed_response
        with patch("editor.agent_service.AGENT_MAX_TOOL_CALLS", 24):
            updated = agent.advance_run(run=run)

        self.assertEqual(updated.status, "failed")
        self.assertIn("tool-call budget", updated.error_message)

    @patch("editor.agent_service._new_openai_client")
    def test_incomplete_without_answer_forces_final_response_instead_of_more_continuations(self, new_client):
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="chat",
            status="in_progress",
            stage="waiting_openai",
            response_id="resp_incomplete",
            tool_calls=[{"source": "biaedge", "name": "search_cases", "type": "mcp_call"}],
        )
        incomplete_response = SimpleNamespace(
            id="resp_incomplete",
            status="incomplete",
            error=None,
            incomplete_details=SimpleNamespace(reason="max_output_tokens"),
            output_text="",
            output=[],
        )

        with patch.object(agent, "_queue_force_final_response", return_value=run) as force_final:
            updated = agent._continue_incomplete_response(run=run, response=incomplete_response)

        self.assertIs(updated, run)
        force_final.assert_called_once_with(run=run, response=incomplete_response)

    @patch("editor.agent_service._new_openai_client")
    def test_advance_run_reassembles_partial_answer_across_continuation(self, new_client):
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="chat",
            status="in_progress",
            stage="waiting_openai",
            response_id="resp_partial",
            metadata=agent._initial_run_metadata(mode="chat", previous_response_id=""),
        )

        partial_response = SimpleNamespace(
            id="resp_partial",
            status="incomplete",
            error=None,
            incomplete_details=SimpleNamespace(reason="max_output_tokens"),
            output_text="Beginning of the answer. ",
            output=[
                SimpleNamespace(
                    type="message",
                    content=[
                        SimpleNamespace(
                            type="output_text",
                            text="Beginning of the answer. ",
                            annotations=[],
                        )
                    ],
                )
            ],
            usage=None,
        )
        queued_response = SimpleNamespace(
            id="resp_tail",
            status="queued",
            error=None,
            usage=None,
            output=[],
        )
        tail_response = SimpleNamespace(
            id="resp_tail",
            status="completed",
            error=None,
            output_text="End of the answer.",
            output=[
                SimpleNamespace(
                    type="message",
                    content=[
                        SimpleNamespace(
                            type="output_text",
                            text="End of the answer.",
                            annotations=[],
                        )
                    ],
                )
            ],
            usage=None,
        )

        def retrieve(response_id, **kwargs):
            if response_id == "resp_partial":
                return partial_response
            if response_id == "resp_tail":
                return tail_response
            raise AssertionError(f"Unexpected response id: {response_id}")

        agent.client.responses.retrieve = retrieve

        with patch.object(agent, "_build_tools", return_value=[]):
            with patch.object(agent, "_create_background_response", return_value=queued_response):
                updated = agent.advance_run(run=run)

        self.assertEqual(updated.response_id, "resp_tail")
        self.assertEqual(updated.status, "queued")

        updated.refresh_from_db()
        completed = agent.advance_run(run=updated)

        self.assertEqual(completed.status, "completed")
        self.assertEqual(
            completed.result_payload["answer"],
            "Beginning of the answer. End of the answer.",
        )

    @patch("editor.agent_service._new_openai_client")
    def test_advance_run_requests_full_text_before_finalizing_quote_answer(self, new_client):
        agent = DocumentResearchAgent(
            document=self.document,
            user=self.user,
        )
        session = DocumentResearchSession.objects.create(document=self.document, user=self.user)
        run = DocumentResearchRun.objects.create(
            session=session,
            mode="chat",
            status="in_progress",
            stage="waiting_openai",
            response_id="resp_policy_search",
            request_payload={
                "message": "Find quotes from case law and quotes from the USCIS Policy Manual for this paragraph.",
                "selected_text": "Selected paragraph text.",
            },
            metadata=agent._initial_run_metadata(mode="chat", previous_response_id=""),
            tool_calls=[
                {
                    "source": "biaedge",
                    "type": "mcp_call",
                    "name": "search_references",
                    "status": "completed",
                    "arguments": {"query": "uscis policy manual change waiver basis", "source_code": "uscis_pm"},
                }
            ],
        )

        completed_response = SimpleNamespace(
            id="resp_policy_search",
            status="completed",
            error=None,
            output_text=(
                "I do not have a verified direct quote from the USCIS Policy Manual in the materials gathered for this turn."
            ),
            output=[
                SimpleNamespace(
                    type="message",
                    content=[
                        SimpleNamespace(
                            type="output_text",
                            text=(
                                "I do not have a verified direct quote from the USCIS Policy Manual in the materials gathered for this turn."
                            ),
                            annotations=[],
                        )
                    ],
                )
            ],
            usage=None,
        )
        queued_response = SimpleNamespace(
            id="resp_quote_verify",
            status="queued",
            error=None,
            usage=None,
            output=[],
        )

        agent.client.responses.retrieve = lambda *args, **kwargs: completed_response

        with patch.object(agent, "_build_tools", return_value=[{"type": "mcp"}]):
            with patch.object(agent, "_create_background_response", return_value=queued_response) as create_background:
                updated = agent.advance_run(run=run)

        self.assertEqual(updated.status, "queued")
        self.assertEqual(updated.stage, "verifying_quotes")
        self.assertEqual(updated.response_id, "resp_quote_verify")
        self.assertTrue(updated.metadata.get("quote_source_verification_attempted"))
        request = create_background.call_args.kwargs
        self.assertEqual(request["tool_choice"], "auto")
        self.assertEqual(request["previous_response_id"], "resp_policy_search")
        self.assertEqual(request["instructions"], agent._run_instructions(run))
        self.assertIn("Turn requirements JSON:", request["input_payload"])
        self.assertIn("missing_full_text_sources: policy, case_law", request["input_payload"])


class OpenAIClientFileServiceTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="openai-file-user", password="secret")
        self.document_type = DocumentType.objects.create(
            name="OpenAI File Cover Letter",
            slug="openai-file-cover-letter",
            category="cover_letter",
            template_content=_sample_tiptap("Template"),
        )
        self.document = Document.objects.create(
            title="OpenAI File Document",
            document_type=self.document_type,
            content=_sample_tiptap("Draft"),
            created_by=self.user,
        )

    def test_sync_client_file_openai_index_uploads_and_indexes_file(self):
        client = SimpleNamespace(
            files=SimpleNamespace(create=lambda **kwargs: SimpleNamespace(id="file_upload_123")),
            vector_stores=SimpleNamespace(
                create=lambda **kwargs: SimpleNamespace(id="vs_client_docs_123"),
                files=SimpleNamespace(
                    create_and_poll=lambda **kwargs: SimpleNamespace(
                        id="vsfile_client_docs_123",
                        status="completed",
                    )
                ),
            ),
        )
        client_file = DocumentClientFile.objects.create(
            document=self.document,
            title="Scanned I-94",
            original_file=SimpleUploadedFile("i94-scan.pdf", b"%PDF-1.4 fake"),
            extracted_text="",
            uploaded_by=self.user,
            metadata={"filename": "i94-scan.pdf", "extension": ".pdf", "text_extracted": False},
        )

        metadata = sync_client_file_openai_index(client_file, client=client)

        client_file.refresh_from_db()
        self.assertEqual(metadata["openai_file_id"], "file_upload_123")
        self.assertEqual(metadata["openai_vector_store_id"], "vs_client_docs_123")
        self.assertEqual(metadata["openai_vector_store_file_id"], "vsfile_client_docs_123")
        self.assertEqual(metadata["openai_index_status"], "completed")
        self.assertTrue(metadata["scan_candidate"])
        self.assertEqual(client_file.metadata["openai_file_id"], "file_upload_123")

    def test_analyze_client_file_with_input_file_sends_original_file_to_responses_api(self):
        recorded_request = {}

        def fake_create(**kwargs):
            recorded_request.update(kwargs)
            return SimpleNamespace(output_text="The I-94 shows admission on March 1, 2024.", output=[])

        client = SimpleNamespace(
            responses=SimpleNamespace(create=fake_create),
        )
        client_file = DocumentClientFile.objects.create(
            document=self.document,
            title="I-94 Record",
            original_file=SimpleUploadedFile("i94-record.pdf", b"%PDF-1.4 fake"),
            extracted_text="",
            uploaded_by=self.user,
            metadata={
                "filename": "i94-record.pdf",
                "extension": ".pdf",
                "openai_file_id": "file_i94_123",
            },
        )

        result = analyze_client_file_with_input_file(
            client_file,
            query="What admission date appears on the I-94?",
            client=client,
        )

        self.assertEqual(result["analysis"], "The I-94 shows admission on March 1, 2024.")
        self.assertEqual(recorded_request["model"], "gpt-5.4")
        self.assertEqual(recorded_request["tool_choice"], "none")
        self.assertEqual(recorded_request["input"][0]["content"][0]["type"], "input_text")
        self.assertEqual(recorded_request["input"][0]["content"][1]["type"], "input_file")
        self.assertEqual(recorded_request["input"][0]["content"][1]["file_id"], "file_i94_123")


class DocumentImportTests(TestCase):
    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        cls._media_root = tempfile.mkdtemp(prefix="editor-import-tests-")
        cls._override = override_settings(MEDIA_ROOT=cls._media_root)
        cls._override.enable()

    @classmethod
    def tearDownClass(cls):
        cls._override.disable()
        shutil.rmtree(cls._media_root, ignore_errors=True)
        super().tearDownClass()

    def setUp(self):
        self.user = User.objects.create_user(username="importer", password="secret")
        self.client.force_login(self.user)
        self.document_type = DocumentType.objects.create(
            name="Imported Brief",
            slug="imported-brief",
            category="brief",
            export_format="court_brief",
            description="Brief imported from Word",
            icon="📘",
        )

    def test_import_docx_to_tiptap_preserves_basic_structure(self):
        content = import_docx_to_tiptap(BytesIO(_build_docx_bytes()))
        nodes = content.get("content", [])

        self.assertEqual(content.get("type"), "doc")
        self.assertEqual(nodes[0]["type"], "heading")
        self.assertEqual(nodes[0]["attrs"]["level"], 1)
        self.assertEqual(nodes[1]["type"], "paragraph")
        self.assertEqual(nodes[2]["type"], "bulletList")
        self.assertEqual(nodes[3]["type"], "table")

    def test_import_docx_package_captures_word_metadata(self):
        package = import_docx_package(BytesIO(_build_docx_bytes()))

        self.assertAlmostEqual(package["metadata"]["page_setup"]["left_margin_pt"], 72.0, places=1)
        paragraph_attrs = package["content"]["content"][1]["attrs"]
        self.assertEqual(paragraph_attrs["word_style"]["name"], "Normal")
        self.assertIn("space_before_pt", paragraph_attrs["paragraph_metrics"])
        first_text_marks = package["content"]["content"][1]["content"][1]["marks"]
        self.assertTrue(any(mark["type"] == "wordRun" for mark in first_text_marks))
        self.assertIn("list_identity", package["content"]["content"][2]["attrs"])

    def test_import_document_creates_editable_document_with_source_docx(self):
        upload = SimpleUploadedFile(
            "existing-brief.docx",
            _build_docx_bytes(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

        response = self.client.post(
            reverse("import_document"),
            {
                "title": "Existing I-751 Brief",
                "document_type": self.document_type.slug,
                "file": upload,
            },
        )

        self.assertEqual(response.status_code, 302)
        document = Document.objects.get(title="Existing I-751 Brief")
        self.assertTrue(bool(document.source_docx))
        self.assertEqual(document.document_type, self.document_type)
        self.assertEqual(document.versions.count(), 1)
        self.assertEqual(document.versions.first().label, "Imported from Word")
        self.assertEqual(document.content.get("type"), "doc")
        self.assertEqual(document.metadata["fidelity_mode"], "proof")
        self.assertEqual(document.metadata["source_docx_info"]["filename"], "existing-brief.docx")

    def test_docx_export_uses_source_docx_template_when_present(self):
        upload = SimpleUploadedFile(
            "existing-brief.docx",
            _build_docx_bytes(margin_inches=1.5),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        document = Document.objects.create(
            title="Imported Export Template",
            document_type=self.document_type,
            content={
                "type": "doc",
                "content": [
                    {"type": "heading", "attrs": {"level": 1}, "content": [{"type": "text", "text": "Updated Heading"}]},
                    {"type": "paragraph", "content": [{"type": "text", "text": "Updated body paragraph."}]},
                ],
            },
            source_docx=upload,
            created_by=self.user,
        )

        response = self.client.get(reverse("export_docx", kwargs={"doc_id": document.id}))

        self.assertEqual(response.status_code, 200)
        exported = DocxDocument(BytesIO(response.content))
        self.assertAlmostEqual(exported.sections[0].left_margin.inches, 1.5, places=2)
        self.assertEqual(exported.paragraphs[0].text, "Updated Heading")


class CoverLetterExportTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="exporter", password="secret")
        self.client.force_login(self.user)
        self.document_type = DocumentType.objects.create(
            name="I-130 Cover Letter",
            slug="i-130-cover-letter-test",
            category="cover_letter",
            export_format="cover_letter",
            template_content=_sample_cover_letter(),
        )
        self.document = Document.objects.create(
            title="I-130 Cover Letter",
            document_type=self.document_type,
            content=_sample_cover_letter(),
            created_by=self.user,
        )

    def test_docx_export_uses_cover_letter_style_anchor(self):
        response = self.client.get(reverse("export_docx", kwargs={"doc_id": self.document.id}))

        self.assertEqual(response.status_code, 200)
        exported = DocxDocument(BytesIO(response.content))
        non_empty = [p.text.strip() for p in exported.paragraphs if p.text.strip()]

        self.assertGreaterEqual(len(non_empty), 10)
        self.assertEqual(non_empty[0], "Chris Hammond Law Firm")
        self.assertIn("Immigration Attorney", non_empty[:6])
        self.assertIn("RE:\tForm I-130, Petition for Alien Relative", non_empty)
        self.assertIn("Respectfully submitted,", non_empty)
        self.assertIn("Christopher Hammond, Esq.", non_empty)
        self.assertGreaterEqual(len(exported.tables), 2)
        self.assertEqual(exported.tables[0].cell(0, 0).text.strip(), "EXHIBIT 1")
        self.assertEqual(exported.tables[0].cell(1, 0).text.strip(), "EXHIBIT 2")


class ProofPreviewViewTests(TestCase):
    def setUp(self):
        self.user = User.objects.create_user(username="proof-user", password="secret")
        self.client.force_login(self.user)
        self.document_type = DocumentType.objects.create(
            name="Proof Brief",
            slug="proof-brief",
            category="brief",
            export_format="court_brief",
            template_content=_sample_tiptap("Template"),
        )
        self.document = Document.objects.create(
            title="Proof Draft",
            document_type=self.document_type,
            content=_sample_tiptap("Proof content."),
            created_by=self.user,
        )

    @patch("editor.views.render_document_proof")
    def test_proof_manifest_endpoint_returns_manifest(self, render_document_proof):
        render_document_proof.return_value = {
            "kind": "document",
            "id": str(self.document.id),
            "preview_available": True,
            "hash": "abc123",
            "pdf_url": "/media/proof.pdf",
            "page_count": 2,
            "pages": [{"index": 1, "image_url": "/media/page-1.png"}],
            "backend": "word_mac",
            "generated_at": "2026-03-16T15:00:00Z",
        }

        response = self.client.post(reverse("proof_refresh", kwargs={"doc_id": self.document.id}))

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["backend"], "word_mac")
        self.document.refresh_from_db()
        self.assertEqual(self.document.metadata["preview_state"]["hash"], "abc123")

    @patch("editor.views.WordRenderService")
    @patch("editor.views.render_document_proof")
    def test_proof_manifest_endpoint_returns_backend_diagnostics_on_error(
        self,
        render_document_proof,
        word_render_service,
    ):
        render_document_proof.side_effect = ProofRenderError("No proof render backend is available.")
        word_render_service.return_value.backend_status.return_value = [
            {
                "name": "soffice",
                "available": False,
                "path": "",
                "detail": "LibreOffice/soffice was not found.",
            }
        ]

        response = self.client.get(reverse("proof_manifest", kwargs={"doc_id": self.document.id}))

        self.assertEqual(response.status_code, 503)
        payload = response.json()
        self.assertEqual(payload["error"], "No proof render backend is available.")
        self.assertEqual(payload["backend_status"][0]["name"], "soffice")

    def test_set_document_style_source_persists_exemplar_override(self):
        exemplar = Exemplar.objects.create(
            title="Style Master",
            document_type=self.document_type,
            kind="style_anchor",
            style_family="uscis_cover_letter",
            original_file=SimpleUploadedFile(
                "style-master.docx",
                _build_docx_bytes(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            created_by=self.user,
        )

        response = self.client.post(
            reverse("document_style_source", kwargs={"doc_id": self.document.id, "exemplar_id": exemplar.id})
        )

        self.assertEqual(response.status_code, 200)
        self.document.refresh_from_db()
        self.assertEqual(self.document.metadata["style_source_exemplar_id"], exemplar.id)
        self.assertEqual(self.document.metadata["style_source_label"], "Style Master")


class ProofRenderBackendTests(TestCase):
    @patch("editor.proof_service.shutil.which", return_value=None)
    def test_soffice_backend_finds_standard_binary_without_path_lookup(self, _which):
        with tempfile.TemporaryDirectory() as tmpdir:
            binary_path = Path(tmpdir) / "soffice"
            binary_path.write_text("#!/bin/sh\nexit 0\n")
            binary_path.chmod(0o755)

            with patch.object(SofficeRenderBackend, "binary_candidates", (str(binary_path),)):
                backend = SofficeRenderBackend()

        self.assertEqual(backend.binary, str(binary_path))
        self.assertTrue(backend.is_available())

    @patch("editor.proof_service._build_pdf_preview_assets", return_value=(1, []))
    @patch("editor.proof_service.tiptap_to_pdf")
    @patch("editor.proof_service.WordRenderService.render_docx_to_pdf")
    def test_render_document_proof_falls_back_to_internal_pdf_when_word_rendering_is_unavailable(
        self,
        render_docx_to_pdf,
        tiptap_to_pdf_mock,
        _build_pdf_preview_assets,
    ):
        render_docx_to_pdf.side_effect = ProofRenderError("No proof render backend is available.")
        tiptap_to_pdf_mock.return_value = BytesIO(b"%PDF-1.4 fallback")
        user = User.objects.create_user(username="proof-fallback-user", password="secret")
        document_type = DocumentType.objects.create(
            name="Fallback Brief",
            slug="fallback-brief",
            category="brief",
            export_format="court_brief",
            template_content=_sample_tiptap("Template"),
        )
        document = Document.objects.create(
            title="Fallback Draft",
            document_type=document_type,
            content=_sample_tiptap("Fallback proof content."),
            created_by=user,
        )

        with tempfile.TemporaryDirectory() as tmpdir:
            with override_settings(MEDIA_ROOT=tmpdir):
                manifest = render_document_proof(document, user=user, force=True)

        self.assertEqual(manifest["backend"], "internal_pdf")
        self.assertFalse(manifest["exact_render"])
        self.assertIn("internal PDF export", manifest["notice"])
        self.assertEqual(manifest["page_count"], 1)


class ExemplarWorkflowTests(TestCase):
    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        cls._media_root = tempfile.mkdtemp(prefix="editor-exemplar-tests-")
        cls._override = override_settings(MEDIA_ROOT=cls._media_root)
        cls._override.enable()

    @classmethod
    def tearDownClass(cls):
        cls._override.disable()
        shutil.rmtree(cls._media_root, ignore_errors=True)
        super().tearDownClass()

    def setUp(self):
        self.user = User.objects.create_user(username="exemplar-user", password="secret")
        self.client.force_login(self.user)
        self.document_type = DocumentType.objects.create(
            name="Exemplar Brief",
            slug="exemplar-brief",
            category="brief",
            export_format="court_brief",
            template_content=_sample_tiptap("Template"),
        )
        self.exemplar = Exemplar.objects.create(
            title="Master Exemplar",
            document_type=self.document_type,
            kind="style_anchor",
            style_family="uscis_cover_letter",
            original_file=SimpleUploadedFile(
                "master-exemplar.docx",
                _build_docx_bytes(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            extracted_text="Master exemplar body text.",
            created_by=self.user,
        )

    def test_open_exemplar_as_draft_creates_proof_mode_document(self):
        response = self.client.post(
            reverse("exemplar_open_as_draft", kwargs={"exemplar_id": self.exemplar.id})
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        created = Document.objects.get(id=payload["document_id"])
        self.assertTrue(bool(created.source_docx))
        self.assertEqual(created.metadata["fidelity_mode"], "proof")
        self.assertEqual(created.metadata["source_exemplar_id"], self.exemplar.id)
        self.assertEqual(created.metadata["style_source_exemplar_id"], self.exemplar.id)

    @patch("editor.exemplar_views.render_exemplar_preview")
    def test_exemplar_preview_endpoint_returns_preview_payload(self, render_exemplar_preview):
        render_exemplar_preview.return_value = {
            "kind": "exemplar",
            "id": self.exemplar.id,
            "preview_available": True,
            "pdf_url": "/media/proof-previews/exemplar.pdf",
            "page_count": 1,
            "pages": [{"index": 1, "image_url": "/media/proof-previews/exemplar-1.png"}],
            "backend": "word_mac",
            "filename": "master-exemplar.docx",
        }

        response = self.client.get(reverse("exemplar_preview", kwargs={"exemplar_id": self.exemplar.id}))

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["page_count"], 1)
        self.assertIn("open_as_draft_url", payload)


class SeedTemplatesTests(TestCase):
    def test_i751_templates_seed_as_three_distinct_cover_letters(self):
        call_command("seed_templates")

        templates = {
            doc_type.slug: doc_type
            for doc_type in DocumentType.objects.filter(slug__icontains="i-751")
        }

        self.assertEqual(
            sorted(templates.keys()),
            [
                "i-751-change-joint-to-waiver-cover-letter",
                "i-751-cover-letter",
                "i-751-waiver-cover-letter",
            ],
        )
        self.assertEqual(templates["i-751-cover-letter"].name, "I-751 Joint Filing Cover Letter")
        self.assertEqual(templates["i-751-waiver-cover-letter"].name, "I-751 Waiver Filing Cover Letter")
        self.assertEqual(
            templates["i-751-change-joint-to-waiver-cover-letter"].name,
            "I-751 Request to Change Joint Filing to Waiver Cover Letter",
        )

        for doc_type in templates.values():
            headings = [
                "".join(
                    child.get("text", "")
                    for child in node.get("content", [])
                    if isinstance(child, dict)
                )
                for node in doc_type.template_content.get("content", [])
                if node.get("type") == "heading"
            ]
            self.assertEqual(headings[:2], ["I. INTRODUCTION", "II. ARGUMENT"])
            self.assertTrue(headings[2].startswith("A."))
            self.assertTrue(headings[3].startswith("B."))
            self.assertTrue(headings[4].startswith("C."))
            self.assertIn("III. CONCLUSION", headings)
